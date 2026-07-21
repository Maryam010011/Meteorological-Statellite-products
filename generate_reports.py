"""
generate_reports.py
===================
Generates for each of the 23 MSG SEVIRI products:
  1. Interpretation Report  -> reports/<slug>.docx
  2. Two-page User Guide    -> guides/<slug>_guide.docx

Products: 12 channels + 11 composites (9 RGB + colorized_ir + natural_color)
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = os.path.dirname(os.path.abspath(__file__))
REP_DIR = os.path.join(BASE, "reports")
GUI_DIR = os.path.join(BASE, "guides")
os.makedirs(REP_DIR, exist_ok=True)
os.makedirs(GUI_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# PRODUCT DEFINITIONS
# ---------------------------------------------------------------------------
PRODUCTS = [
    # (key, display_name, wavelength, type, description, physical, colors,
    #  agri, avia, nrm, ndm)
    ("VIS006","VIS 0.6 µm","0.6 µm","channel",
     "Visible broadband channel measuring reflected solar radiation.",
     "Reflectance (%). Measures sunlight reflected from clouds, land, sea. "
     "Bright = highly reflective (thick cloud, snow). Dark = low reflectance (ocean, vegetation).",
     "Grayscale: white=cloud/snow, light grey=thin cloud, dark=clear land/ocean.",
     "Crop monitoring: cloud cover assessment for irrigation scheduling; detect hail-damaged crop fields (bright scars). "
     "Vegetation mapping using reflectance contrast between crops and bare soil.",
     "Real-time cloud cover for flight planning and in-flight hazard avoidance. "
     "Identify convective cloud tops, fog banks, and dust layers that threaten operations.",
     "Land use / vegetation mapping; snow cover extent for water resource assessment; "
     "deforestation monitoring by tracking reflectance changes over time.",
     "Track wildfire smoke plumes; detect large dust storms; identify flood extents "
     "when water bodies expand; monitor volcanic ash clouds."
    ),
    ("VIS008","VIS 0.8 µm","0.8 µm","channel",
     "Near-infrared visible channel sensitive to vegetation and land surface.",
     "Reflectance (%). Vegetation strongly reflects at 0.8 µm (red-edge). "
     "Useful for distinguishing vegetation from bare soil and cloud.",
     "Grayscale: vegetated areas appear lighter than in VIS006 due to NIR reflectance boost.",
     "Vegetation health index proxy; distinguish crops from fallow land; "
     "monitor growing season onset and senescence.",
     "Distinguish low cloud/fog from land features; better land contrast than VIS006 "
     "for situational awareness near airports.",
     "Vegetation density mapping; agricultural land classification; "
     "forest vs grassland discrimination.",
     "Forest fire burned area mapping; vegetation stress from drought visible "
     "as reduced NIR reflectance."
    ),
    ("IR_016","NIR 1.6 µm","1.6 µm","channel",
     "Near-infrared channel sensitive to ice/water phase discrimination.",
     "Reflectance (%). Ice clouds appear dark (low reflectance at 1.6 µm), "
     "water clouds appear bright. Snow/ice on surface also dark.",
     "Grayscale: water clouds = bright, ice clouds = dark, snow/ice surface = dark.",
     "Distinguish frozen precipitation from liquid — critical for frost/freeze warnings. "
     "Snow cover mapping for melt-water runoff forecasting.",
     "Identify ice-phase clouds (icing risk for aircraft) vs. liquid water clouds. "
     "Key channel for icing severity nowcasting.",
     "Snow/ice extent mapping; glacier monitoring; differentiate water from ice in lakes.",
     "Detect ice storms; map snow accumulation from blizzards; "
     "identify hail-producing cloud systems (ice-phase tops)."
    ),
    ("IR_039","IR 3.9 µm","3.9 µm","channel",
     "Mid-infrared channel sensitive to both emitted thermal and reflected solar radiation.",
     "Brightness temperature (K). At 3.9 µm both emission and solar reflection contribute. "
     "Hot spots (fires, volcanoes) appear very bright. Used for fire detection and fog.",
     "Gray_r: cold = bright (cloud tops), warm = dark. Fire pixels appear as very bright hot spots.",
     "Active fire and burned area detection; soil moisture proxy at night; "
     "frost detection (cold surface = bright).",
     "Fog detection (critical for runway visual range); volcanic activity monitoring "
     "near flight routes.",
     "Volcanic hot spot monitoring; active fire mapping for forest management; "
     "thermal anomaly detection.",
     "Real-time wildfire detection and spread; volcanic eruption hot spots; "
     "industrial fire monitoring."
    ),
    ("WV_062","WV 6.2 µm","6.2 µm","channel",
     "Upper-tropospheric water vapour channel (600-300 hPa).",
     "Brightness temperature (K). Senses water vapour in upper troposphere. "
     "Bright (cold) = high moisture or high cloud. Dark (warm) = dry upper troposphere.",
     "Gray_r: bright = moist upper troposphere / upper-level clouds, dark = dry subsidence.",
     "Large-scale rainfall forecasting; upper-level moisture for crop water stress; "
     "track monsoon onset and retreat.",
     "Upper-level jet stream identification; turbulence forecasting in clear air; "
     "identify regions of deep convection.",
     "Drought monitoring via persistent warm (dry) anomalies; "
     "track inter-tropical convergence zone.",
     "Detect intense cyclone outflow and upper-level divergence; "
     "track tornado-producing supercell environment."
    ),
    ("WV_073","WV 7.3 µm","7.3 µm","channel",
     "Mid-tropospheric water vapour channel (500-300 hPa).",
     "Brightness temperature (K). Senses mid-tropospheric moisture. "
     "Bridges gap between surface and upper-level WV channels.",
     "Gray_r: bright = moist mid-levels, dark = dry.",
     "Mid-level moisture for rainfall probability; track moisture transport into crop regions.",
     "Mid-level turbulence diagnosis; identify areas of conditional instability.",
     "Evapotranspiration monitoring; inter-basin moisture flux analysis.",
     "Identify low-level jet and moisture convergence feeding severe storms; "
     "tropical cyclone structure analysis."
    ),
    ("IR_087","IR 8.7 µm","8.7 µm","channel",
     "Infrared window channel sensitive to dust and volcanic aerosols.",
     "Brightness temperature (K). Absorbs in silicate mineral dust absorption band. "
     "Used in dust RGB composites; sensitive to thin cirrus and SO2.",
     "Gray_r: cold cloud = bright, warm surface = dark; dust appears anomalously cold.",
     "Dust storm impact on agriculture; track dust deposition events; "
     "soil erosion source region identification.",
     "Dust and volcanic ash detection for flight safety (SIGMET/SIGMET support); "
     "Saharan dust over Mediterranean and Middle East.",
     "Desertification monitoring; dust transport pathways from arid regions; "
     "volcanic SO2 cloud tracking.",
     "Dust storm early warning; volcanic ash cloud tracking after eruption; "
     "sandstorm impact assessment."
    ),
    ("IR_097","IR 9.7 µm","9.7 µm","channel",
     "Ozone absorption band channel.",
     "Brightness temperature (K). Sensitive to ozone in stratosphere. "
     "Mainly used for total column ozone estimation and stratospheric analysis.",
     "Gray_r: provides total column ozone proxy; low ozone = warmer BT.",
     "UV radiation exposure linked to ozone; indirect crop damage assessment.",
     "Ozone hole monitoring relevant to UV radiation forecast for aviation routes.",
     "Total ozone column monitoring; UV index forecasting for ecosystem health.",
     "Stratospheric intrusion events that can affect surface weather systems."
    ),
    ("IR_108","IR 10.8 µm","10.8 µm","channel",
     "Primary thermal infrared window channel — the workhorse of SEVIRI.",
     "Brightness temperature (K). Near-transparent to atmosphere. "
     "Measures cloud top or surface temperature. Cold = high thick cloud, warm = clear warm surface.",
     "Gray_r: bright white = cold high cloud, mid-grey = mid-level cloud/warm surface, dark = hot desert/ocean.",
     "Cloud cover and temperature for crop microclimate; detect cold air outbreaks damaging crops; "
     "estimate Land Surface Temperature (LST) for evapotranspiration models.",
     "Primary channel for cloud top temperature analysis, SIGMET issuance, "
     "CB detection, and Aviation Meteorological Forecast (TAF/METAR support).",
     "Sea Surface Temperature (SST) estimation; LST for drought indices; "
     "urban heat island monitoring; thermal anomaly detection.",
     "Track tropical cyclone intensity via warm core and cold cloud tops; "
     "flash flood nowcasting from deep convection; volcanic eruption plume temperature."
    ),
    ("IR_120","IR 12.0 µm","12.0 µm","channel",
     "Infrared split-window channel used alongside IR_108 for SST and dust.",
     "Brightness temperature (K). Slightly more water-vapour-absorbing than IR_108. "
     "Paired with IR_108 (split-window technique) for SST and dust detection.",
     "Gray_r: appearance similar to IR_108; differences from IR_108 reveal moisture/dust.",
     "Split-window SST for fishery and irrigation water temperature; "
     "soil moisture proxy via split-window emissivity.",
     "Fog discrimination from low cloud; complement to IR_108 for cloud classification.",
     "SST mapping for marine resource management; "
     "desertification monitoring via surface emissivity changes.",
     "Dust storm detection via IR_108 − IR_120 difference; "
     "wildfire smoke plume tracking."
    ),
    ("IR_134","IR 13.4 µm","13.4 µm","channel",
     "CO2 absorption band channel for cloud height estimation.",
     "Brightness temperature (K). CO2 absorption reduces weighting to surface. "
     "Used with other IR channels for cloud-top height (CTH) retrieval.",
     "Gray_r: senses mid-to-upper troposphere temperature; CTH derivation channel.",
     "Cloud height for precipitation type (rain vs snow boundary); "
     "frost risk assessment from radiative cooling.",
     "Cloud top height for aircraft route planning; separation from high clouds; "
     "SIGMET cloud height estimates.",
     "Cloud climatology for solar energy resource assessment; "
     "tropical cloud-top height monitoring.",
     "Deep convection height — direct proxy for hail and severe wind potential; "
     "tropical cyclone cloud-top cooling rate for rapid intensification."
    ),
    ("HRV","HRV 0.7 µm (High Res)","0.7 µm","channel",
     "High-Resolution Visible channel at ~1 km spatial resolution.",
     "Reflectance (%). Panchromatic visible channel with ~3x finer resolution than standard SEVIRI. "
     "Shows fine-scale cloud structure, coastlines, and urban features.",
     "Grayscale: finest spatial detail — cloud streets, city boundaries, "
     "small-scale convection, river/lake edges visible.",
     "Field-level cloud assessment; fine-scale fog mapping for agricultural frost risk; "
     "high-resolution crop damage mapping after storms.",
     "Fine cloud detail for terminal area forecasts; fog/stratus in valleys; "
     "convective initiation on small scales.",
     "Coastline change detection; urban expansion; fine-scale deforestation; "
     "river course changes.",
     "High-res cyclone eye and eyewall structure; "
     "micro-scale convective cell tracking; fine flood boundary mapping."
    ),
    ("natural_color","Natural Color RGB","0.6/0.8/1.6 µm","composite",
     "True-colour-like composite: R=VIS006, G=VIS008, B=IR_016. "
     "Daylight only. Gives near-natural appearance.",
     "RGB image. Land appears green/brown, ocean blue-black, clouds white. "
     "Ice/snow appears cyan (dark blue in 1.6 µm makes blue channel low).",
     "Green = vegetated land, brown = bare soil/desert, white = thick water cloud, "
     "cyan/blue = ice cloud or snow.",
     "Crop greenness assessment; land cover change; vegetation phenology; "
     "snow cover mapping for melt-water runoff.",
     "Intuitive cloud/land/ocean discrimination; convective cell identification; "
     "volcanic plume color distinction.",
     "Land use / land cover mapping; forest health; seasonal vegetation cycle; "
     "coastal sedimentation.",
     "Wildfire smoke (grey-brown); dust (beige); flood extent over farmland; "
     "volcanic ash (grey-brown cloud)."
    ),
    ("airmass","Airmass RGB","WV/IR","composite",
     "Airmass RGB: R=WV6.2−WV7.3, G=IR9.7−IR10.8, B=WV6.2. "
     "Identifies air mass types and jet streams.",
     "RGB highlights dynamic tropopause, ozone, and moisture. "
     "Distinct colors identify warm/cold air masses and jet stream location.",
     "Red/orange = warm dry descending air (subtropical ridge). "
     "Green = moist tropical air. Blue = cold polar air. Yellow = jet stream region.",
     "Identify air mass type governing growing season weather; "
     "cold air outbreak detection for freeze warnings.",
     "Jet stream position for flight routing and turbulence avoidance; "
     "rapid cyclogenesis indicators.",
     "Large-scale moisture transport for water resource planning; "
     "drought-favorable air mass identification.",
     "Explosive cyclogenesis (bomb cyclone) detection; "
     "identify dry intrusion fueling wildfire spread; tornado environment diagnosis."
    ),
    ("dust","Dust RGB","IR 8.7/10.8/12.0 µm","composite",
     "Dust RGB: R=IR12.0−IR10.8, G=IR10.8−IR8.7, B=IR10.8. "
     "Specifically designed to detect mineral dust.",
     "Exploits differential absorption of dust in mid-IR channels. "
     "Dust appears pink/magenta. Clouds appear white/cyan. Clear warm surface appears dark red.",
     "Pink/magenta = active dust plume or transported dust layer. "
     "White = thick cloud. Cyan = thin cloud. Dark red/brown = hot clear surface.",
     "Dust deposition damage to crops; identify dust source regions; "
     "desertification monitoring; soil erosion events.",
     "Critical for aviation safety — dust reduces visibility below minimums; "
     "SIGMET-level dust storm identification over Middle East, Africa, Arabian Sea.",
     "Desertification extent; dust source region mapping; "
     "sand dune migration; soil degradation.",
     "Haboob (dust storm wall) tracking; large-scale Saharan and Arabian dust events; "
     "dust-induced visibility hazards for transport and human health."
    ),
    ("ash","Ash RGB","IR 8.7/10.8/12.0 µm","composite",
     "Ash RGB: R=IR12.0−IR10.8, G=IR10.8−IR8.7, B=IR10.8. "
     "Tuned for volcanic ash detection.",
     "Similar channel combination to Dust RGB but volcanic ash has distinct "
     "spectral signature. Ash appears red/orange, clouds white, clear sky dark.",
     "Red/orange = volcanic ash cloud. White = meteorological cloud. "
     "Dark = clear sky. Purple/pink = mixed ash+cloud.",
     "Ash fallout on farmland can fertilize or destroy crops; "
     "identify affected agricultural zones post-eruption.",
     "Highest priority aviation product — volcanic ash destroys jet engines. "
     "Supports Volcanic Ash Advisory Centre (VAAC) operations.",
     "Volcanic ash hazard zone mapping; post-eruption land surface change; "
     "lava flow detection when combined with IR_039.",
     "Volcanic eruption hazard mapping; ash cloud trajectory tracking; "
     "evacuation zone delineation."
    ),
    ("convection","Convection RGB","WV/VIS/IR","composite",
     "Convection RGB: R=WV6.2−WV7.3, G=IR3.9−IR10.8, B=NIR1.6−VIS0.6. "
     "Highlights rapidly developing convection.",
     "Identifies vigorous growing convective cells through channel differences "
     "sensitive to deep convective updrafts and overshooting tops.",
     "Yellow/green = active convective cell. Orange = intense convection. "
     "Blue = cold anvil. White = very deep convection.",
     "Severe thunderstorm and hail threat to crops; "
     "convective rainfall for irrigation management; harvest weather windows.",
     "Rapid CB (Cumulonimbus) detection for inflight avoidance; "
     "SIGMET issuance support; nowcasting for departure/arrival.",
     "Convective rainfall mapping for soil recharge; "
     "identify flood-prone catchments.",
     "Flash flood nowcasting; tornado-producing supercell identification; "
     "hail storm tracking; severe wind event nowcasting."
    ),
    ("fog","Fog RGB","IR 3.9/10.8/12.0 µm","composite",
     "Fog/Stratus RGB: R=IR12.0−IR3.9, G=IR10.8−IR3.9, B=IR10.8. "
     "Night-time fog and low stratus detection.",
     "Works day and night. Fog and low stratus appear orange/yellow "
     "due to difference between 3.9 µm emissivity and longer IR channels.",
     "Orange/yellow = fog or low stratus. Dark = clear sky. "
     "Cyan/white = medium/high cloud. Black/dark blue = warm clear ocean.",
     "Frost and fog impact on harvesting and field work; "
     "low cloud persistence forecasts for agricultural planning.",
     "Fog and low stratus at airports — primary product for runway visual range; "
     "CAT I/II/III ILS decision support.",
     "Fog frequency mapping for tourism/transport infrastructure planning; "
     "valley fog delineation.",
     "Radiation fog events that cause road accidents and disruption; "
     "sea fog impact on coastal shipping."
    ),
    ("night_microphysics","Night Microphysics RGB","IR 3.9/10.8/12.0 µm","composite",
     "Night Microphysics RGB: R=IR12.0−IR10.8, G=IR10.8−IR3.9, B=IR10.8. "
     "Cloud phase and particle size discrimination at night.",
     "Works at night. Distinguishes ice from water clouds and thin from thick cloud. "
     "Particle size determines icing risk and precipitation type.",
     "Green = water cloud (small droplets — fog/stratus). "
     "Red = thin ice cloud. White = thick ice cloud. Dark = clear sky.",
     "Night fog and frost risk for sensitive crops; "
     "detect ice-phase precipitation overnight for harvest protection.",
     "Night-time icing in cloud for aircraft; "
     "discrimination of freezing rain from snow; "
     "nocturnal convection over land.",
     "Night cloud cover for solar energy resource planning; "
     "maritime fog and low cloud monitoring.",
     "Night-time severe storm tracking; nocturnal MCS (Mesoscale Convective System); "
     "winter storm ice vs snow discrimination."
    ),
    ("day_severe_storms","Day Severe Storms RGB","VIS/NIR/WV","composite",
     "Day Severe Storms RGB: R=WV6.2−WV7.3, G=IR3.9−IR10.8, B=NIR1.6−VIS0.6. "
     "Identifies severe convective storm potential by day.",
     "Highlights overshooting tops, strong updrafts, and cloud phase in "
     "severe convective environments. Day-time only.",
     "Yellow/orange = severe storm with overshooting top. "
     "Green = active convection. Blue = cold anvil. White = deep convective core.",
     "Hail threat mapping for crop insurance; "
     "severe storm track for pre-harvest decision making.",
     "Inflight severe turbulence and icing avoidance; "
     "CB nowcasting for terminal area management.",
     "Convective rainfall for flash flood assessment; "
     "wind damage mapping in forested areas.",
     "Tornado and supercell tracking; hail storm path; "
     "derecho (wind storm) identification; flash flood nowcasting."
    ),
    ("colorized_ir_clouds","IR 10.8 µm Colorized","10.8 µm","composite",
     "Colorized version of IR 10.8 µm using a rainbow color table. "
     "Enhances cloud top temperature differences visually.",
     "Same physical data as IR_108 but displayed with colorized enhancement. "
     "Cold cloud tops appear deep blue/purple, warm surfaces red/brown.",
     "Deep blue/purple = very cold high cloud tops (CB, deep convection, cyclones). "
     "Yellow/green = mid-level cloud. Red/brown = warm clear surface.",
     "Rainfall probability from cold cloud tops; "
     "freeze alert from cold surface colors; "
     "cloud cover persistence for crop planning.",
     "Rapid identification of deep convection for flight briefing; "
     "cyclone intensity from cloud top temperature.",
     "SST anomaly visualization; land surface temperature patterns; "
     "thermal plume detection from power plants or upwelling.",
     "Tropical cyclone intensity and track; "
     "MCS rainfall estimation; "
     "volcanic plume height estimation from temperature."
    ),
]

# ---------------------------------------------------------------------------
# DOCX HELPERS
# ---------------------------------------------------------------------------

def set_heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if color:
        for run in h.runs:
            run.font.color.rgb = RGBColor(*color)
    return h

def add_para(doc, text, bold=False, italic=False, size=10, color=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = RGBColor(*color)
    return p

def add_table_row(table, cells, bold_first=True):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = text
        if bold_first and i == 0:
            for run in cell.paragraphs[0].runs:
                run.bold = True

def shade_cell(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)

def page_break(doc):
    doc.add_page_break()

def set_margins(doc, top=2, bottom=2, left=2, right=2):
    for section in doc.sections:
        section.top_margin    = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin   = Cm(left)
        section.right_margin  = Cm(right)

# ---------------------------------------------------------------------------
# REPORT GENERATOR
# ---------------------------------------------------------------------------

def make_report(p):
    key, name, wl, ptype, desc, phys, colors, agri, avia, nrm, ndm = p
    sl = key.lower().replace(" ","_").replace(".","p").replace("/","_")

    doc = Document()
    set_margins(doc)

    # Title block
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("MSG SEVIRI IODC — Meteorological Product Interpretation Report")
    r.bold = True; r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0x00, 0x38, 0x6B)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub.add_run("{} ({})".format(name, wl))
    r2.bold = True; r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor(0x00, 0x70, 0xC0)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Satellite: Meteosat-9 IODC  |  Sub-satellite: 45.5°E  |  "
                 "Instrument: SEVIRI  |  Product type: {}".format(ptype.capitalize())).font.size = Pt(9)
    doc.add_paragraph()

    # 1. Product Overview
    set_heading(doc, "1.  Product Overview", 1, (0,56,107))
    add_para(doc, desc, size=10)
    doc.add_paragraph()

    # 2. Physical Basis
    set_heading(doc, "2.  Physical Basis and Calibration", 1, (0,56,107))
    add_para(doc, phys, size=10)
    doc.add_paragraph()

    # 3. Image Interpretation
    set_heading(doc, "3.  Image Interpretation Guide", 1, (0,56,107))
    add_para(doc, "Color / Tone Key:", bold=True, size=10)
    add_para(doc, colors, size=10)
    doc.add_paragraph()

    # 4. Applications table
    set_heading(doc, "4.  Sector Applications", 1, (0,56,107))

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    hdr = tbl.rows[0].cells
    hdr[0].text = "Sector"; hdr[1].text = "Application Details"
    for cell in hdr:
        shade_cell(cell, "003869")
        for p_ in cell.paragraphs:
            for r in p_.runs:
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.bold = True

    for sector, text in [
        ("Agriculture",              agri),
        ("Aviation",                 avia),
        ("Natural Resource Monitoring", nrm),
        ("Natural Disaster Monitoring", ndm),
    ]:
        row = tbl.add_row().cells
        row[0].text = sector
        row[1].text = text
        shade_cell(row[0], "D6E4F0")
        for p_ in row[0].paragraphs:
            for r in p_.runs: r.bold = True

    doc.add_paragraph()

    # 5. Limitations
    set_heading(doc, "5.  Known Limitations", 1, (0,56,107))
    limits = []
    if ptype == "channel" and wl.startswith("0"):
        limits.append("Daylight only — no data at night (reflected solar channels).")
    if "WV" in key:
        limits.append("Cannot retrieve surface information — senses free-tropospheric moisture only.")
    if ptype == "composite":
        limits.append("Composite colours are product-specific — refer to colour key before interpretation.")
    if key in ("natural_color","day_severe_storms"):
        limits.append("Daylight hours only — product is dark or invalid at night.")
    limits += [
        "15-minute repeat cycle limits rapid-change phenomena detection between scans.",
        "Spatial resolution 3 km at nadir (1 km HRV) — sub-grid features not resolved.",
        "Edge-of-disk distortion at high satellite zenith angles (> 70°).",
    ]
    for lim in limits:
        doc.add_paragraph(lim, style="List Bullet")

    doc.add_paragraph()

    # 6. References
    set_heading(doc, "6.  References", 1, (0,56,107))
    refs = [
        "EUMETSAT. (2024). MSG SEVIRI Level 1.5 Product User Manual. EUM/MSG/ICD/105.",
        "Schmetz, J. et al. (2002). An Introduction to Meteosat Second Generation. BAMS 83(7).",
        "EUMETSAT. (2023). RGB Composite Quick Guides — {} Product.".format(name),
        "WMO. (2018). Manual on the Global Observing System. WMO-No. 544.",
        "Kidder, S. & Vonder Haar, T. (1995). Satellite Meteorology. Academic Press.",
    ]
    for ref in refs:
        doc.add_paragraph(ref, style="List Bullet")

    path = os.path.join(REP_DIR, "{}_report.docx".format(sl))
    doc.save(path)
    print("  Report saved: {}".format(os.path.basename(path)))
    return path


# ---------------------------------------------------------------------------
# USER GUIDE GENERATOR  (2 pages)
# ---------------------------------------------------------------------------

def make_guide(p):
    key, name, wl, ptype, desc, phys, colors, agri, avia, nrm, ndm = p
    sl = key.lower().replace(" ","_").replace(".","p").replace("/","_")

    doc = Document()
    set_margins(doc, top=1.8, bottom=1.8, left=2, right=2)

    # ---- PAGE 1 ----
    # Header bar
    hdr_p = doc.add_paragraph()
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr_p.add_run("USER GUIDE  |  MSG SEVIRI IODC  |  Meteosat-9  |  45.5°E")
    r.bold = True; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(255,255,255)
    shading = hdr_p._p.get_or_add_pPr()

    # Title
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = t.add_run("{}\n{}".format(name, wl))
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0x00, 0x38, 0x6B)

    doc.add_paragraph()

    # Product overview box as table
    ov_tbl = doc.add_table(rows=1, cols=1)
    ov_tbl.style = "Table Grid"
    cell = ov_tbl.rows[0].cells[0]
    shade_cell(cell, "EBF3FB")
    cell.text = ""
    cell.paragraphs[0].add_run("PRODUCT OVERVIEW").bold = True
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    cell.add_paragraph(desc).runs[0].font.size = Pt(9) if cell.paragraphs else None
    p2 = cell.add_paragraph()
    p2.add_run("Physical basis: ").bold = True
    p2.add_run(phys[:180] + ("..." if len(phys)>180 else "")).font.size = Pt(9)
    doc.add_paragraph()

    # HOW TO READ THIS IMAGE
    set_heading(doc, "How to Read This Image", 2, (0,56,107))
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    h1, h2 = tbl2.rows[0].cells
    h1.text = "What You See"; h2.text = "What It Means"
    shade_cell(h1,"003869"); shade_cell(h2,"003869")
    for c in [h1,h2]:
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True
                rr.font.color.rgb = RGBColor(255,255,255)

    # Parse color key into rows
    color_items = [x.strip() for x in colors.replace(". ","\n").split("\n") if x.strip() and "=" in x]
    if not color_items:
        color_items = [colors]
    for item in color_items:
        if "=" in item:
            left, right = item.split("=",1)
        else:
            left, right = item, ""
        row = tbl2.add_row().cells
        row[0].text = left.strip()
        row[1].text = right.strip()
        shade_cell(row[0],"D6E4F0")

    doc.add_paragraph()

    # Quick-Reference Sector Table
    set_heading(doc, "Sector Quick-Reference", 2, (0,56,107))
    tbl3 = doc.add_table(rows=1, cols=3)
    tbl3.style = "Table Grid"
    c0,c1,c2 = tbl3.rows[0].cells
    for c,txt in [(c0,"Sector"),(c1,"Key Use"),(c2,"Watch For")]:
        c.text = txt; shade_cell(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)

    sector_data = [
        ("Agriculture",    agri[:90]+"...",  agri[90:150]+"..."),
        ("Aviation",       avia[:90]+"...",  avia[90:150]+"..."),
        ("Resource Mgmt",  nrm[:90]+"...",   nrm[90:150]+"..."),
        ("Disaster Mgmt",  ndm[:90]+"...",   ndm[90:150]+"..."),
    ]
    for sec, use, watch in sector_data:
        row = tbl3.add_row().cells
        row[0].text = sec; row[1].text = use; row[2].text = watch
        shade_cell(row[0],"D6E4F0")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold = True

    # ---- PAGE 2 ----
    page_break(doc)

    set_heading(doc, "Detailed Applications", 1, (0,56,107))

    for sector, text in [
        ("Agriculture", agri),
        ("Aviation", avia),
        ("Natural Resource Monitoring", nrm),
        ("Natural Disaster Monitoring", ndm),
    ]:
        set_heading(doc, sector, 2, (0,112,192))
        add_para(doc, text, size=10)
        doc.add_paragraph()

    # Do's and Don'ts
    set_heading(doc, "Do's and Don'ts", 1, (0,56,107))
    dos_donts = doc.add_table(rows=1, cols=2)
    dos_donts.style = "Table Grid"
    dh, dnh = dos_donts.rows[0].cells
    dh.text = "DO"; dnh.text = "DON'T"
    shade_cell(dh,"006400"); shade_cell(dnh,"8B0000")
    for c,col in [(dh,RGBColor(255,255,255)),(dnh,RGBColor(255,255,255))]:
        for pp in c.paragraphs:
            for rr in pp.runs: rr.bold=True; rr.font.color.rgb=col

    dos = [
        "Cross-reference with numerical weather model output.",
        "Use multiple channels/composites for confident interpretation.",
        "Account for solar zenith angle in visible channel interpretation.",
        "Check sensing time stamp before operational use.",
    ]
    donts = [
        "Do not use VIS channels for night-time analysis.",
        "Do not confuse dust/ash signatures without multi-channel check.",
        "Do not assume cloud-free pixels mean no fog at surface.",
        "Do not use edge-of-disk pixels for quantitative analysis.",
    ]
    for do, dont in zip(dos, donts):
        row = dos_donts.add_row().cells
        row[0].text = do; row[1].text = dont

    doc.add_paragraph()

    # Footer
    ft = doc.add_paragraph()
    ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ft = ft.add_run(
        "Generated from Meteosat-9 SEVIRI Level 1.5 data  |  "
        "NASTAP Assignment 4  |  Product: {}".format(name))
    r_ft.font.size = Pt(8)
    r_ft.font.color.rgb = RGBColor(128,128,128)
    r_ft.italic = True

    path = os.path.join(GUI_DIR, "{}_guide.docx".format(sl))
    doc.save(path)
    print("  Guide  saved: {}".format(os.path.basename(path)))
    return path


# ---------------------------------------------------------------------------
# MAIN
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    print("Generating {} reports and {} guides ...".format(len(PRODUCTS), len(PRODUCTS)))
    print("\n--- REPORTS ---")
    for p in PRODUCTS:
        try:
            make_report(p)
        except Exception as e:
            print("  ERROR report {}: {}".format(p[0], e))

    print("\n--- GUIDES ---")
    for p in PRODUCTS:
        try:
            make_guide(p)
        except Exception as e:
            print("  ERROR guide {}: {}".format(p[0], e))

    print("\nDone.")
    print("Reports -> {}".format(REP_DIR))
    print("Guides  -> {}".format(GUI_DIR))
