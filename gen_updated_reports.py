"""
gen_updated_reports.py
======================
Regenerates all 21 interpretation reports and 21 user guides,
updated to reflect the 40-color composite variants now available.
Overwrites reports/ and guides/ folders.
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE    = os.path.dirname(os.path.abspath(__file__))
REP_DIR = os.path.join(BASE, "reports")
GUI_DIR = os.path.join(BASE, "guides")
os.makedirs(REP_DIR, exist_ok=True)
os.makedirs(GUI_DIR, exist_ok=True)

N_COLORS = 40   # number of color variants per composite

def shade(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:fill"),hex_color); shd.set(qn("w:val"),"clear")
    tcPr.append(shd)

def margins(doc,t=2,b=2,l=2,r=2):
    for s in doc.sections:
        s.top_margin=Cm(t); s.bottom_margin=Cm(b)
        s.left_margin=Cm(l); s.right_margin=Cm(r)

def h1(doc,txt,color=(0,56,107)):
    h=doc.add_heading(txt,level=1)
    for r in h.runs: r.font.color.rgb=RGBColor(*color)
    return h

def h2(doc,txt,color=(0,112,192)):
    h=doc.add_heading(txt,level=2)
    for r in h.runs: r.font.color.rgb=RGBColor(*color)
    return h

def para(doc,txt,sz=10.5,bold=False,italic=False,color=None,center=False):
    pp=doc.add_paragraph()
    if center: pp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=pp.add_run(txt)
    r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*color)
    return pp

def bullet(doc,txt,sz=10):
    pp=doc.add_paragraph(style="List Bullet")
    r=pp.add_run(txt); r.font.size=Pt(sz)
    return pp

def two_col(doc,rows,headers=None,hdr_bg="003869"):
    tbl=doc.add_table(rows=0,cols=2)
    tbl.style="Table Grid"
    if headers:
        hr=tbl.add_row().cells
        hr[0].text=headers[0]; hr[1].text=headers[1]
        for c in hr:
            shade(c,hdr_bg)
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for l,r in rows:
        row=tbl.add_row().cells
        row[0].text=str(l); row[1].text=str(r)
        shade(row[0],"EBF3FB")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True
    return tbl

def slug(name):
    return name.lower().replace(" ","_").replace("/","_").replace(".","p").replace("(","").replace(")","")

# ── Product database ──────────────────────────────────────────────────────────
PRODUCTS = [
    ("VIS006","VIS 0.6 um","0.6 um","channel",
     "Visible broadband — reflected solar radiation. Daytime only.",
     "Reflectance (%). Bright=cloud/snow. Dark=ocean/vegetation.",
     "White=thick cloud. Grey=thin cloud/haze. Dark=clear surface.",
     "Crop cloud cover; hail scar detection; irrigation scheduling.",
     "Cloud cover for flight planning; convective cell identification.",
     "Land cover mapping; snow extent for water resource planning.",
     "Wildfire smoke; dust storms; flood extent; volcanic ash."),
    ("VIS008","VIS 0.8 um","0.8 um","channel",
     "Near-infrared visible — enhanced vegetation contrast. Daytime only.",
     "Reflectance (%). Vegetation reflects strongly at 0.8 um.",
     "Lighter vegetation than VIS006. Better land/cloud contrast.",
     "Crop health; growing season monitoring; land classification.",
     "Low cloud vs land discrimination near airports.",
     "Vegetation density; agricultural land classification.",
     "Burned area mapping; vegetation stress from drought."),
    ("IR_016","NIR 1.6 um","1.6 um","channel",
     "Near-infrared — ice vs water cloud discrimination. Daytime only.",
     "Reflectance (%). Ice clouds dark. Water clouds bright. Snow dark.",
     "Bright=water cloud. Dark=ice cloud or snow surface.",
     "Freeze/frost warnings; snow cover for melt-water forecasting.",
     "Aircraft icing risk — ice vs liquid cloud discrimination.",
     "Snow/ice extent; glacier monitoring.",
     "Ice storms; hail-producing cloud (ice-phase tops)."),
    ("IR_039","IR 3.9 um","3.9 um","channel",
     "Mid-infrared — fire detection, fog, mixed emission/reflection.",
     "Brightness temperature (K). Hot spots very bright. Day+night.",
     "Gray_r: bright=cold cloud. Very bright pixels=fire hot spots.",
     "Active fire detection; soil moisture; frost detection.",
     "Fog/runway visual range; volcanic activity near routes.",
     "Active fire mapping; thermal anomaly detection.",
     "Real-time wildfire detection; volcanic eruption hot spots."),
    ("WV_062","WV 6.2 um","6.2 um","channel",
     "Upper-tropospheric water vapour (600-300 hPa). Day and night.",
     "Brightness temperature (K). Bright=moist/high cloud. Dark=dry.",
     "Bright=moist upper troposphere. Dark=dry subsiding air.",
     "Large-scale rainfall forecasting; monsoon onset/retreat.",
     "Upper-level jet stream; clear-air turbulence forecasting.",
     "Drought monitoring; ITCZ tracking.",
     "Cyclone outflow; tornado environment diagnosis."),
    ("WV_073","WV 7.3 um","7.3 um","channel",
     "Mid-tropospheric water vapour (500-300 hPa). Day and night.",
     "Brightness temperature (K). Bridges surface and upper WV.",
     "Bright=moist mid-levels. Dark=dry subsidence.",
     "Mid-level moisture for rainfall probability.",
     "Mid-level turbulence; conditional instability.",
     "Evapotranspiration monitoring; inter-basin moisture flux.",
     "Low-level jet; moisture convergence for severe storms."),
    ("IR_087","IR 8.7 um","8.7 um","channel",
     "Silicate mineral dust absorption band. Day and night.",
     "Brightness temperature (K). Dust anomalously cold vs surface.",
     "Gray_r: dust appears as anomalously bright (cold) patches.",
     "Dust damage to crops; desertification; soil erosion.",
     "Dust/ash detection for SIGMET; Middle East/Africa routes.",
     "Desertification extent; dust source mapping.",
     "Haboob tracking; Saharan/Arabian dust events."),
    ("IR_097","IR 9.7 um","9.7 um","channel",
     "Stratospheric ozone absorption band. Day and night.",
     "Brightness temperature (K). Sensitive to column ozone.",
     "Gray_r: ozone-dependent BT. Low ozone = warmer appearance.",
     "UV radiation exposure; crop damage assessment.",
     "Ozone hole monitoring for polar routes.",
     "Total ozone column; UV index for ecosystems.",
     "Stratospheric intrusion affecting surface weather."),
    ("IR_108","IR 10.8 um","10.8 um","channel",
     "Primary thermal window — cloud/surface temperature. Day and night.",
     "Brightness temperature (K). Cold=high cloud. Warm=clear surface.",
     "Gray_r: bright white=deep cloud. Dark=hot desert/ocean.",
     "Cloud temperature for crop microclimate; freeze alerts; LST.",
     "Primary CB detection; SIGMET support; TAF/METAR.",
     "SST; LST for drought; urban heat island.",
     "Tropical cyclone intensity; flash flood nowcasting."),
    ("IR_120","IR 12.0 um","12.0 um","channel",
     "Split-window channel for SST and dust retrieval. Day and night.",
     "Brightness temperature (K). Paired with IR_108 for split-window.",
     "Gray_r: similar to IR_108. Differences reveal moisture/dust.",
     "Split-window SST for fisheries; soil moisture proxy.",
     "Fog discrimination from low cloud.",
     "SST mapping; desertification via surface emissivity.",
     "Dust detection (IR108-IR120 difference); smoke tracking."),
    ("IR_134","IR 13.4 um","13.4 um","channel",
     "CO2 absorption — cloud-top height retrieval. Day and night.",
     "Brightness temperature (K). CO2 limits surface sensing.",
     "Gray_r: mid-to-upper troposphere temperature; CTH channel.",
     "Cloud height for precipitation type; frost risk.",
     "Cloud-top height for separation/routing; SIGMET CTH.",
     "Cloud climatology for solar energy; tropical CTH.",
     "Deep convection height; hail/severe wind potential."),
    ("HRV","HRV 0.7 um","0.7 um","channel",
     "High-Resolution Visible (~1 km) — fine detail. Daytime only.",
     "Reflectance (%). Panchromatic. 3x finer than standard SEVIRI.",
     "Finest detail: cloud streets, city edges, small convection.",
     "Field-level cloud; fine fog for frost risk; storm damage.",
     "Fine cloud for terminal area forecasts; valley fog.",
     "Coastline change; urban expansion; deforestation.",
     "Cyclone eye/eyewall; micro-scale convection; flood boundary."),
    ("natural_color","Natural Color","0.6/0.8/1.6 um","composite",
     "R=VIS006 G=VIS008 B=IR016. True-colour-like. Daytime only.",
     "RGB. Land green/brown. Ocean blue-black. Cloud white. Snow cyan.",
     "Green=vegetation. White=water cloud. Cyan=ice cloud/snow.",
     "Crop greenness; land cover change; snow mapping.",
     "Intuitive cloud/land/ocean view; convective cells.",
     "Land use mapping; forest health; coastal sedimentation.",
     "Wildfire smoke; dust plumes; flood extent over farmland."),
    ("airmass","Airmass RGB","WV/IR","composite",
     "R=WV062-WV073 G=IR097-IR108 B=WV062. Air mass identification.",
     "RGB. Colors identify warm/cold air masses and jet stream.",
     "Red/orange=warm dry descending. Green=moist tropical. Blue=polar.",
     "Air mass type for growing season; cold outbreak freeze warnings.",
     "Jet stream for routing; rapid cyclogenesis indicators.",
     "Large-scale moisture transport; drought-favorable air mass.",
     "Bomb cyclone detection; dry intrusion fueling wildfires."),
    ("dust","Dust RGB","IR 8.7/10.8/12.0 um","composite",
     "R=IR120-IR108 G=IR108-IR087 B=IR108. Mineral dust detection.",
     "RGB. Dust=pink/magenta. Cloud=white/cyan. Surface=dark red.",
     "Pink/magenta=dust plume. White=thick cloud. Dark=clear hot surface.",
     "Dust crop damage; desertification; soil erosion events.",
     "Critical aviation product; SIGMET dust; Middle East/Africa.",
     "Desertification; dust source mapping; sand dune migration.",
     "Haboob tracking; Saharan/Arabian dust; visibility hazards."),
    ("ash","Ash RGB","IR 8.7/10.8/12.0 um","composite",
     "R=IR120-IR108 G=IR108-IR087 B=IR108. Volcanic ash detection.",
     "RGB. Ash=red/orange. Cloud=white. Clear=dark.",
     "Red/orange=volcanic ash. White=meteorological cloud. Dark=clear.",
     "Ash fallout on farmland; affected agricultural zones.",
     "Highest priority aviation — ash destroys engines; VAAC support.",
     "Volcanic hazard mapping; lava flow with IR039.",
     "Eruption hazard; ash trajectory; evacuation delineation."),
    ("convection","Convection RGB","WV/VIS/IR","composite",
     "R=WV062-WV073 G=IR039-IR108 B=NIR-VIS. Rapid convection. Day.",
     "RGB. Yellow/green=active cell. Orange=intense. Blue=cold anvil.",
     "Yellow=convective cell. Orange=intense. Blue=anvil. White=deep.",
     "Severe storm/hail threat; convective rainfall; harvest windows.",
     "CB detection for avoidance; SIGMET; nowcasting.",
     "Convective rainfall; flood-prone catchments.",
     "Flash flood nowcasting; supercell; hail; severe wind."),
    ("fog","Fog RGB","IR 3.9/10.8/12.0 um","composite",
     "R=IR120-IR039 G=IR108-IR039 B=IR108. Fog/stratus. Day+night.",
     "RGB. Fog=orange/yellow. Clear=dark. High cloud=cyan/white.",
     "Orange/yellow=fog or low stratus. Dark=clear. Cyan=medium cloud.",
     "Frost/fog impact on harvesting; low cloud persistence.",
     "Primary airport fog product; CAT I/II/III ILS support.",
     "Fog frequency mapping; valley fog delineation.",
     "Radiation fog causing road accidents; sea fog for shipping."),
    ("night_microphysics","Night Microphysics RGB","IR 3.9/10.8/12.0 um","composite",
     "R=IR120-IR108 G=IR108-IR039 B=IR108. Cloud phase at night.",
     "RGB. Night-time ice vs water cloud discrimination.",
     "Green=water cloud (fog). Red=thin ice. White=thick ice. Dark=clear.",
     "Night fog/frost for crops; ice precipitation overnight.",
     "Night icing in cloud; freezing rain vs snow.",
     "Night cloud for solar energy; maritime fog.",
     "Nocturnal MCS; winter storm ice vs snow discrimination."),
    ("day_severe_storms","Day Severe Storms RGB","VIS/NIR/WV","composite",
     "R=WV062-WV073 G=IR039-IR108 B=NIR-VIS. Severe storms. Day only.",
     "RGB. Yellow=overshooting top. Green=active convection. Blue=anvil.",
     "Yellow/orange=severe storm. Green=active. Blue=anvil. White=core.",
     "Hail threat to crops; severe storm for pre-harvest decisions.",
     "Inflight severe turbulence; CB nowcasting.",
     "Convective rainfall; wind damage in forests.",
     "Tornado/supercell; hail storm; derecho; flash flood."),
    ("colorized_ir_clouds","IR 10.8 Colorized","10.8 um","composite",
     "IR_108 with rainbow color table. Enhanced cloud top temperatures.",
     "Same as IR_108 but displayed with colorized enhancement.",
     "Deep blue=very cold high cloud (CB/cyclone). Red=warm surface.",
     "Rainfall probability; freeze alert; cloud persistence.",
     "Deep convection for flight briefing; cyclone intensity.",
     "SST anomaly; LST patterns; thermal plume detection.",
     "Tropical cyclone intensity; MCS rainfall; volcanic plume height."),
]

# ── Report generator ──────────────────────────────────────────────────────────
def make_report(p):
    key,name,wl,ptype,desc,phys,colors,agri,avia,nrm,ndm = p
    sl = slug(key)
    doc = Document(); margins(doc)

    # Title
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run("MSG SEVIRI IODC — Meteorological Product Interpretation Report")
    r.bold=True;r.font.size=Pt(14);r.font.color.rgb=RGBColor(0,56,107)
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r2=sp.add_run(f"{name} ({wl})")
    r2.bold=True;r2.font.size=Pt(13);r2.font.color.rgb=RGBColor(0,112,192)
    mp=doc.add_paragraph(); mp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    mp.add_run(f"Satellite: Meteosat-9 IODC  |  Sub-satellite: 45.5E  |  "
               f"Type: {ptype.capitalize()}  |  "
               f"Images per scene: {'1' if ptype=='channel' else f'1 standard + {N_COLORS} color variants'}").font.size=Pt(9)
    doc.add_paragraph()

    h1(doc,"1.  Product Overview")
    para(doc,desc)
    doc.add_paragraph()

    h1(doc,"2.  Physical Basis")
    para(doc,phys)
    doc.add_paragraph()

    h1(doc,"3.  Image Interpretation — Color Key")
    para(doc,colors)
    doc.add_paragraph()

    # Color variants info for composites
    if ptype=="composite":
        h1(doc,"3b.  40-Color Variant Composites")
        para(doc,
            f"In addition to the standard EUMETSAT-calibrated RGB version, "
            f"this composite is also rendered in {N_COLORS} alternative color palettes. "
            f"Each version uses the same underlying satellite data converted to a "
            f"luminance channel, then displayed with a different matplotlib colormap. "
            f"This provides {N_COLORS} visual perspectives on the same physical scene, "
            f"making different meteorological features more visible depending on the palette chosen.")
        doc.add_paragraph()
        two_col(doc,[
            ("Standard composite","Satpy EUMETSAT-calibrated RGB — scientifically correct colors"),
            ("01-Gray_r","Cold features bright — standard IR-like view"),
            ("03-Jet / Rainbow","Full spectrum — general overview"),
            ("05-Inferno","Cold dark, hot bright — thermal emphasis"),
            ("08-Viridis","Perceptually uniform — accessible to colorblind users"),
            ("10-Spectral","Multi-feature discrimination"),
            ("14-Coolwarm","Diverging — anomaly detection"),
            ("17-YlOrRd_r","Surface temperature emphasis"),
            ("28-Ocean","Marine feature enhancement"),
            (f"... +{N_COLORS-9} more","Total: {N_COLORS} color variants per composite per scene"),
        ], headers=["Color Variant","Best Use"])
        doc.add_paragraph()

    h1(doc,"4.  Sector Applications")
    tbl=doc.add_table(rows=1,cols=2); tbl.style="Table Grid"
    hr=tbl.rows[0].cells
    hr[0].text="Sector"; hr[1].text="Application Details"
    for c in hr:
        shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for sector,txt in [("Agriculture",agri),("Aviation",avia),
                        ("Natural Resource Monitoring",nrm),("Natural Disaster Monitoring",ndm)]:
        row=tbl.add_row().cells
        row[0].text=sector; row[1].text=txt
        shade(row[0],"D6E4F0")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True
    doc.add_paragraph()

    h1(doc,"5.  Output Files for This Product")
    out_rows=[("Product type",ptype.capitalize()),("Scenes covered","22 unique scenes"),]
    if ptype=="channel":
        out_rows+=[("PNG per scene","1 (standard grayscale/IR with colorbar)"),
                   ("GeoTIFF per scene","1 (georeferenced, LZW-compressed)"),
                   ("Total PNGs","22"),("Total GeoTIFFs","22")]
    else:
        out_rows+=[("Standard PNG per scene","1 (EUMETSAT calibrated RGB)"),
                   (f"{N_COLORS}-color PNGs per scene",f"{N_COLORS} (geostationary projection, circular disk)"),
                   ("GeoTIFF per scene","1 (standard composite, georeferenced)"),
                   ("Total standard PNGs","22"),
                   (f"Total {N_COLORS}-color PNGs",str(22*N_COLORS)),
                   ("Total GeoTIFFs","22")]
    two_col(doc,out_rows,headers=["Item","Value"])
    doc.add_paragraph()

    h1(doc,"6.  Known Limitations")
    limits=[]
    if wl.startswith("0") or key in ("HRV","natural_color","day_severe_storms","convection"):
        limits.append("Daytime only — visible/NIR channels require solar illumination.")
    if ptype=="composite":
        limits.append("Standard composite colors are EUMETSAT-calibrated — do not apply custom stretch.")
        limits.append(f"The {N_COLORS} color variants use luminance conversion — RGB color meaning differs from standard composite.")
    limits+=["15-minute repeat cycle — rapid sub-scan-period events may be missed.",
             "3 km spatial resolution (1 km HRV) — sub-grid features not resolved.",
             "Edge-of-disk distortion at high satellite zenith angles (>70 degrees)."]
    for lim in limits: bullet(doc,lim)
    doc.add_paragraph()

    h1(doc,"7.  References")
    for ref in ["EUMETSAT. (2024). MSG SEVIRI Level 1.5 Product User Manual. EUM/MSG/ICD/105.",
                "Schmetz, J. et al. (2002). An Introduction to Meteosat Second Generation. BAMS 83(7).",
                f"EUMETSAT. (2023). RGB Composite Quick Guide — {name}.",
                "WMO. (2018). Manual on the Global Observing System. WMO-No. 544."]:
        bullet(doc,ref)

    path=os.path.join(REP_DIR,f"{sl}_report.docx")
    doc.save(path)
    return path


# ── Guide generator ───────────────────────────────────────────────────────────
def make_guide(p):
    key,name,wl,ptype,desc,phys,colors,agri,avia,nrm,ndm = p
    sl = slug(key)
    doc = Document(); margins(doc,t=1.8,b=1.8,l=2,r=2)

    # PAGE 1
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=tp.add_run(f"USER GUIDE  |  {name}  |  {wl}  |  Meteosat-9 IODC 45.5E")
    r.bold=True;r.font.size=Pt(14);r.font.color.rgb=RGBColor(0,56,107)
    doc.add_paragraph()

    # Overview box
    ov=doc.add_table(rows=1,cols=1); ov.style="Table Grid"
    oc=ov.rows[0].cells[0]; shade(oc,"EBF3FB")
    oc.paragraphs[0].add_run("PRODUCT OVERVIEW").bold=True
    oc.paragraphs[0].runs[0].font.size=Pt(10)
    ab=oc.add_paragraph(desc); [setattr(r.font,"size",Pt(9)) for r in ab.runs]
    pb=oc.add_paragraph(); pb.add_run("Physical basis: ").bold=True
    pb.add_run(phys[:150]+("..." if len(phys)>150 else "")).font.size=Pt(9)
    if ptype=="composite":
        pc=oc.add_paragraph()
        pc.add_run(f"Color variants: ").bold=True
        pc.add_run(f"Available in {N_COLORS} color palettes in addition to standard RGB.").font.size=Pt(9)
    doc.add_paragraph()

    h2(doc,"How to Read This Image")
    tbl2=doc.add_table(rows=1,cols=2); tbl2.style="Table Grid"
    h1c,h2c=tbl2.rows[0].cells
    h1c.text="What You See"; h2c.text="What It Means"
    for c in [h1c,h2c]:
        shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for item in [x.strip() for x in colors.split(".") if x.strip() and "=" in x]:
        if "=" in item:
            left,right=item.split("=",1)
            row=tbl2.add_row().cells
            row[0].text=left.strip(); row[1].text=right.strip()
            shade(row[0],"D6E4F0")
    doc.add_paragraph()

    h2(doc,"Quick-Reference: Sectors")
    tbl3=doc.add_table(rows=1,cols=2); tbl3.style="Table Grid"
    c0,c1=tbl3.rows[0].cells
    for c,t in [(c0,"Sector"),(c1,"Key Application")]:
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for sec,txt in [("Agriculture",agri[:100]+"..."),("Aviation",avia[:100]+"..."),
                    ("Resource Monitoring",nrm[:100]+"..."),("Disaster Monitoring",ndm[:100]+"...")]:
        row=tbl3.add_row().cells
        row[0].text=sec; row[1].text=txt
        shade(row[0],"D6E4F0")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True

    # PAGE 2
    doc.add_page_break()
    h2(doc,"Detailed Applications")
    for sec,txt in [("Agriculture",agri),("Aviation",avia),
                    ("Natural Resource Monitoring",nrm),("Natural Disaster Monitoring",ndm)]:
        h2(doc,sec,(0,112,192))
        para(doc,txt,sz=10)
        doc.add_paragraph()

    if ptype=="composite":
        h2(doc,f"Using the {N_COLORS} Color Variants",(0,112,192))
        para(doc,
            f"This composite is available in {N_COLORS} color palettes. "
            f"Each uses identical underlying data. Choose based on the feature you want to highlight:",sz=10)
        two_col(doc,[
            ("Gray / Gray_r",    "Classic IR-style — familiar to operational forecasters"),
            ("Jet / Rainbow",    "Full spectrum — general overview, presentations"),
            ("Inferno / Plasma", "Thermal emphasis — fire, cloud top temperature"),
            ("Viridis / Cividis","Colorblind-accessible — scientific publications"),
            ("Spectral",         "Multi-feature — complex scenes with many phenomena"),
            ("Coolwarm / Seismic","Diverging — anomaly detection around a reference"),
            ("Ocean / Terrain",  "Geographic context — surface and marine features"),
            ("Hot / Copper",     "Hot spot emphasis — fire, volcanic, industrial"),
        ], headers=["Palette Group","Best Application"])
        doc.add_paragraph()

    h2(doc,"Do's and Don'ts")
    tbl4=doc.add_table(rows=1,cols=2); tbl4.style="Table Grid"
    dh,dnh=tbl4.rows[0].cells
    dh.text="DO"; dnh.text="DON'T"
    shade(dh,"006400"); shade(dnh,"8B0000")
    for c,col in [(dh,RGBColor(255,255,255)),(dnh,RGBColor(255,255,255))]:
        for pp in c.paragraphs:
            for rr in pp.runs: rr.bold=True; rr.font.color.rgb=col
    dos=["Cross-reference with NWP model output.",
         "Use multiple channels/composites together.",
         "Check sensing timestamp before operational use.",
         f"Use color variants to highlight different features."]
    donts=["Do not use VIS channels at night.",
           "Do not re-stretch EUMETSAT composite colors.",
           "Do not confuse dust/ash without multi-channel check.",
           "Do not use edge-of-disk pixels for quantitative analysis."]
    for do,dont in zip(dos,donts):
        row=tbl4.add_row().cells
        row[0].text=do; row[1].text=dont

    doc.add_paragraph()
    ft=doc.add_paragraph(); ft.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=ft.add_run(f"NASTAP Assignment 4  |  {name}  |  Meteosat-9 IODC 45.5E  |  "
                 f"22 scenes  |  {1 if ptype=='channel' else 1+N_COLORS} images/scene")
    r.italic=True; r.font.size=Pt(8); r.font.color.rgb=RGBColor(128,128,128)

    path=os.path.join(GUI_DIR,f"{sl}_guide.docx")
    doc.save(path)
    return path


# ── Main ──────────────────────────────────────────────────────────────────────
if __name__=="__main__":
    print(f"Generating {len(PRODUCTS)} updated reports ...")
    for prod in PRODUCTS:
        try:
            p=make_report(prod); print(f"  Report: {os.path.basename(p)}")
        except Exception as e: print(f"  ERROR report {prod[0]}: {e}")

    print(f"\nGenerating {len(PRODUCTS)} updated guides ...")
    for prod in PRODUCTS:
        try:
            p=make_guide(prod); print(f"  Guide:  {os.path.basename(p)}")
        except Exception as e: print(f"  ERROR guide {prod[0]}: {e}")

    print(f"\nDone. Reports -> {REP_DIR}")
    print(f"      Guides  -> {GUI_DIR}")
