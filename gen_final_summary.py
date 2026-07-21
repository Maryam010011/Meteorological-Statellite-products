"""
gen_final_summary.py
====================
Generates three documents after the full batch is complete:
  1. per_file_summary.docx  - per .nat file image count breakdown
  2. updated_workflow.docx  - updated workflow with 40-color composites
  3. updated_reports.docx   - updated product list and counts
"""
import os, json, glob
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE     = os.path.dirname(os.path.abspath(__file__))
OUT_ROOT = os.path.join(BASE, "output_v2")

# ── helpers ──────────────────────────────────────────────────────────────────
def shade(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement("w:shd")
    shd.set(qn("w:fill"),hex_color); shd.set(qn("w:val"),"clear")
    tcPr.append(shd)

def margins(doc,t=2,b=2,l=2.5,r=2):
    for s in doc.sections:
        s.top_margin=Cm(t);s.bottom_margin=Cm(b)
        s.left_margin=Cm(l);s.right_margin=Cm(r)

def h1(doc,txt,color=(0,56,107)):
    h=doc.add_heading(txt,level=1)
    for r in h.runs: r.font.color.rgb=RGBColor(*color)
    return h

def h2(doc,txt,color=(0,112,192)):
    h=doc.add_heading(txt,level=2)
    for r in h.runs: r.font.color.rgb=RGBColor(*color)
    return h

def p(doc,txt,sz=10.5,bold=False,italic=False,color=None,center=False):
    para=doc.add_paragraph()
    if center: para.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=para.add_run(txt)
    r.bold=bold;r.italic=italic;r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*color)
    return para

def tbl_row(tbl, cells, bg=None, bold_first=True):
    row=tbl.add_row().cells
    for i,(c,v) in enumerate(zip(row,cells)):
        c.text=str(v)
        if bg: shade(c,bg)
        if bold_first and i==0:
            for pp in c.paragraphs:
                for rr in pp.runs: rr.bold=True
    return row

def hdr_row(tbl, labels, bg="003869"):
    row=tbl.add_row().cells
    for c,lbl in zip(row,labels):
        c.text=lbl; shade(c,bg)
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True
                rr.font.color.rgb=RGBColor(255,255,255)
    return row

# ── Data collection ───────────────────────────────────────────────────────────
def collect_stats():
    """Read output_v2 folder and count actual files per scene."""
    summary_path = os.path.join(OUT_ROOT, "summary.json")
    with open(summary_path, encoding="utf-8") as f:
        summary = json.load(f)

    scenes = []
    for m in summary["manifests"]:
        ts  = m.get("sensing_start","")[:19]
        end = m.get("sensing_end","")[:19]
        fname   = m.get("file","")
        platform= m.get("platform","Meteosat-9")

        # Build folder name
        try:
            dt = datetime.strptime(ts[:16], "%Y-%m-%d %H:%M")
            folder = dt.strftime("%Y%m%dT%H%MZ")
        except:
            folder = ""

        scene_dir = os.path.join(OUT_ROOT, folder)

        # Count actual files
        ch_pngs  = len(glob.glob(os.path.join(scene_dir,"channels","*.png")))
        co_pngs  = len(glob.glob(os.path.join(scene_dir,"composites","*.png")))
        tifs     = len(glob.glob(os.path.join(scene_dir,"geotiff","*.tif")))

        # Count 40-color composites
        comp40_dir = os.path.join(scene_dir, "composites_40")
        comp40_pngs = 0
        comp40_types = 0
        if os.path.exists(comp40_dir):
            for sub in os.listdir(comp40_dir):
                sub_path = os.path.join(comp40_dir, sub)
                if os.path.isdir(sub_path):
                    cnt = len(glob.glob(os.path.join(sub_path,"*.png")))
                    comp40_pngs += cnt
                    if cnt > 0: comp40_types += 1

        total = ch_pngs + co_pngs + comp40_pngs

        scenes.append({
            "file":     fname,
            "start":    ts,
            "end":      end,
            "folder":   folder,
            "platform": platform,
            "channels": ch_pngs,
            "composites": co_pngs,
            "comp40_types": comp40_types,
            "comp40_pngs": comp40_pngs,
            "geotiffs": tifs,
            "total_pngs": ch_pngs + co_pngs + comp40_pngs,
            "total_files": ch_pngs + co_pngs + comp40_pngs + tifs,
        })
    return scenes


# ── DOCUMENT 1: Per-file summary ──────────────────────────────────────────────
def make_per_file_summary(scenes):
    doc = Document()
    margins(doc)

    # Cover
    p(doc,"MSG SEVIRI IODC — Complete Image Generation Summary",
      sz=18,bold=True,color=(0,56,107),center=True)
    p(doc,"Per .nat File Breakdown: Channels + Composites + 40-Color Variants",
      sz=12,italic=True,color=(80,80,80),center=True)
    p(doc,f"Generated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}  |  "
         f"Meteosat-9 IODC 45.5E  |  22 unique scenes",
      sz=9,italic=True,color=(120,120,120),center=True)
    doc.add_paragraph()

    # Grand totals box
    total_ch  = sum(s["channels"]    for s in scenes)
    total_co  = sum(s["composites"]  for s in scenes)
    total_40  = sum(s["comp40_pngs"] for s in scenes)
    total_tif = sum(s["geotiffs"]    for s in scenes)
    total_png = total_ch + total_co + total_40
    total_all = total_png + total_tif

    gt = doc.add_table(rows=1,cols=6)
    gt.style="Table Grid"
    hdr_row(gt,["Channel PNGs","Standard Composites","40-Color Composites",
                "Total PNGs","GeoTIFFs","Grand Total Files"],"003869")
    tbl_row(gt,[total_ch,total_co,total_40,total_png,total_tif,total_all],
            bg="D6E4F0",bold_first=False)
    doc.add_paragraph()

    # Per-file table
    h1(doc,"Per .nat File — Image Count Breakdown")
    p(doc,"Every row shows one .nat file and exactly how many images it produced.",sz=10)
    doc.add_paragraph()

    # Group by date
    date_groups = {}
    for s in scenes:
        date = s["start"][:10]
        date_groups.setdefault(date,[]).append(s)

    date_colors = {
        "2026-06-22":"FFFFFF",
        "2026-06-23":"FFF2CC",
        "2026-06-29":"E2EFDA",
        "2026-07-06":"FCE4D6",
    }

    main_tbl = doc.add_table(rows=0,cols=9)
    main_tbl.style="Table Grid"
    hdr_row(main_tbl,[
        "File","Scan Start (UTC)","Scan End (UTC)",
        "Channels\n(12)","Composites\n(9)","40-Color\nComposites",
        "Total\nPNGs","GeoTIFFs","Total\nFiles"],"003869")

    for date, group in date_groups.items():
        bg = date_colors.get(date,"FFFFFF")

        # Date header row
        row = main_tbl.add_row().cells
        # Merge all cells for date header
        row[0].text = f"DATE: {date}"
        shade(row[0],"C5D9F1")
        for pp in row[0].paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.size=Pt(10)
                rr.font.color.rgb=RGBColor(0,56,107)
        for i in range(1,9):
            shade(row[i],"C5D9F1")

        # Scene rows
        date_ch=date_co=date_40=date_tif=0
        for s in group:
            tbl_row(main_tbl,[
                s["file"],
                s["start"][11:16]+" UTC",
                s["end"][11:16]+" UTC",
                s["channels"],
                s["composites"],
                f"{s['comp40_pngs']}\n({s['comp40_types']} types×40)",
                s["total_pngs"],
                s["geotiffs"],
                s["total_files"],
            ], bg=bg, bold_first=True)
            date_ch +=s["channels"]; date_co+=s["composites"]
            date_40 +=s["comp40_pngs"]; date_tif+=s["geotiffs"]

        # Date subtotal
        dt=date_ch+date_co+date_40
        tbl_row(main_tbl,[
            f"  Subtotal ({len(group)} scenes)","","",
            date_ch,date_co,date_40,dt,date_tif,dt+date_tif
        ],bg="E2EFDA",bold_first=True)

    doc.add_paragraph()

    # Explanation
    h2(doc,"What Each Column Means")
    rows_exp=[
        ("Channels (12)",
         "12 individual SEVIRI channel images per scene: VIS006, VIS008, NIR 1.6, "
         "IR 3.9, WV 6.2, WV 7.3, IR 8.7, IR 9.7, IR 10.8, IR 12.0, IR 13.4, HRV. "
         "Each saved as PNG + GeoTIFF."),
        ("Composites (9)",
         "9 standard EUMETSAT RGB composites: Natural Color, Airmass, Dust, Ash, "
         "Convection, Fog, Night Microphysics, Day Severe Storms, IR Colorized. "
         "Standard Satpy enhancement colors."),
        ("40-Color Composites",
         "Each of the 9 composites rendered in 40 different color palettes "
         "(gray, rainbow, jet, inferno, plasma, viridis, spectral, coolwarm, etc.). "
         "9 types × 40 colors = 360 images per scene."),
        ("Total PNGs",
         "12 channels + 9 standard composites + 360 color-variant composites = 381 PNG per scene."),
        ("GeoTIFFs",
         "21 GeoTIFF files per scene (12 channels + 9 composites), "
         "LZW-compressed, georeferenced in eqc projection."),
    ]
    exp_tbl=doc.add_table(rows=0,cols=2)
    exp_tbl.style="Table Grid"
    hdr_row(exp_tbl,["Column","Explanation"])
    for lft,rgt in rows_exp:
        tbl_row(exp_tbl,[lft,rgt],bg="EBF3FB")

    path=os.path.join(BASE,"per_file_summary.docx")
    doc.save(path)
    print(f"Saved: per_file_summary.docx")
    return path

# ── DOCUMENT 2: Updated Workflow ──────────────────────────────────────────────
def make_updated_workflow(scenes):
    doc = Document()
    margins(doc)

    p(doc,"MSG SEVIRI IODC — Updated Complete Workflow",
      sz=20,bold=True,color=(0,56,107),center=True)
    p(doc,"Including 40-Color Composite Generation",
      sz=13,italic=True,color=(0,112,192),center=True)
    doc.add_paragraph()

    h1(doc,"1.  Complete Pipeline Overview")
    p(doc,"Each .nat file passes through the following pipeline to produce all images:",sz=10)
    doc.add_paragraph()

    steps=[
        ("INPUT","One .nat file","Binary EUMETSAT MSG SEVIRI Level 1.5 native file, "
         "~11.8 MB, containing 12 calibrated spectral channels + header metadata."),
        ("STEP 1","Hard-link rename","Temporary EUMETSAT-format filename created via os.link(). "
         "Original file never modified. Deleted in finally: block after processing."),
        ("STEP 2","Satpy Scene load","Scene(reader='seviri_l1b_native') opens the file. "
         "Reads header: platform, sensing_start, sensing_end, sub-satellite longitude."),
        ("STEP 3","12 Channel images","Each channel loaded, resampled to 2048×2048 eqc grid, "
         "rendered as grayscale/IR PNG with colorbar + metadata label. "
         "Saved as PNG + GeoTIFF."),
        ("STEP 4","9 Standard composites","Each composite loaded on native scene, "
         "Satpy's EUMETSAT-calibrated get_enhanced_image() applied, "
         "scipy.ndimage.zoom to 2048×2048, rendered as RGB PNG. Saved as PNG + GeoTIFF."),
        ("STEP 5","40-Color composites","Each of the 9 composites converted to luminance, "
         "then rendered 40 times with different matplotlib colormaps "
         "(gray, jet, inferno, plasma, viridis, spectral, coolwarm, etc.). "
         "Geostationary projection ensures coastlines stay inside circular disk. "
         "Saved as PNG only (360 images per scene)."),
        ("STEP 6","Manifest","manifest.json + manifest.csv written recording every "
         "product path, status, and sensing timestamp."),
        ("OUTPUT","Per-scene total",
         "12 channel PNGs + 9 composite PNGs + 360 color-variant PNGs + 21 GeoTIFFs "
         "= 381 PNGs + 21 TIFs = 402 files per scene."),
    ]

    flow_tbl=doc.add_table(rows=0,cols=3)
    flow_tbl.style="Table Grid"
    hdr_row(flow_tbl,["Stage","Name","Description"])
    colors=["003869","375623","833C00","7030A0","0070C0","375623","003869","C00000"]
    for (stage,name,desc),col in zip(steps,colors):
        row=flow_tbl.add_row().cells
        row[0].text=stage; row[1].text=name; row[2].text=desc
        shade(row[0],col)
        for pp in row[0].paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
        shade(row[1],"EBF3FB")
        for pp in row[1].paragraphs:
            for rr in pp.runs: rr.bold=True
    doc.add_paragraph()

    h1(doc,"2.  The 40 Color Palettes")
    p(doc,"The same composite data is rendered 40 times, each with a different "
      "matplotlib colormap. This gives meteorologists 40 visual perspectives "
      "on the same physical data — different palettes highlight different features.",sz=10)
    doc.add_paragraph()

    palettes=[
        ("01-Gray_r","Standard IR gray — cold bright"),
        ("02-Gray","Warm bright gray"),("03-Jet","Rainbow jet"),
        ("04-HSV","HSV spectrum"),("05-Inferno","Thermal inferno"),
        ("06-Magma","Magma"),("07-Plasma","Plasma"),("08-Viridis","Viridis"),
        ("09-Cividis","Cividis"),("10-Spectral","Spectral"),
        ("11-Spectral_r","Spectral reversed"),("12-RdYlBu","Red-Yellow-Blue"),
        ("13-RdYlBu_r","Red-Yellow-Blue inverted"),("14-Coolwarm","Cool-Warm"),
        ("15-BWR","Blue-White-Red"),("16-Seismic","Seismic"),
        ("17-YlOrRd_r","Yellow-Orange-Red"),("18-YlOrRd","YlOrRd inverted"),
        ("19-Hot","Hot"),("20-Hot_r","Hot reversed"),
        ("21-AFMhot","AFM Hot"),("22-Gist_heat","Gist Heat"),
        ("23-Copper","Copper"),("24-Autumn","Autumn"),("25-Summer","Summer"),
        ("26-Spring","Spring"),("27-Winter","Winter"),("28-Ocean","Ocean"),
        ("29-Terrain","Terrain"),("30-Gist_earth","Gist Earth"),
        ("31-NCAR","NCAR"),("32-Gist_rainbow","Gist Rainbow"),
        ("33-NIPY","NIPY Spectral"),("34-Turbo","Turbo"),
        ("35-GNUPlot","GNUPlot"),("36-GNUPlot2","GNUPlot2"),
        ("37-CMR","CMR Map"),("38-CubeHelix","Cube Helix"),
        ("39-BRG","BRG"),("40-Tab20b","Tab 20B"),
    ]

    pal_tbl=doc.add_table(rows=0,cols=4)
    pal_tbl.style="Table Grid"
    hdr_row(pal_tbl,["#","Colormap","Description","Best For"])
    bests=["Night IR, cloud tops","Colorful overview","Cloud temperatures",
           "Full spectrum display","Cold cloud detection","Smooth gradients",
           "Perceptually uniform","Accessible/colorblind","Accessible warm",
           "Multi-feature display","Reversed spectral","Diverging temperature",
           "Warm emphasis","Diverging anomalies","Anomaly detection",
           "Diverging strong","Surface temperature","Cool emphasis",
           "Fire/heat detection","Hot inverted","Hot spots","Thermal gradients",
           "Warm tones","Seasonal transitions","Vegetation contrast",
           "Cold emphasis","Ice/snow","Ocean features","Terrain mapping",
           "Earth-like tones","Atmospheric levels","Full color range",
           "Science standard","Fast overview","Classic science",
           "High contrast","Full spectrum","Smooth diverging",
           "Color reversed","Categorical"]

    bg_cycle=["EBF3FB","E2EFDA","FCE4D6","EAE0F0","FFF2CC"]
    for i,(num_name,desc) in enumerate(palettes):
        num,name=num_name.split("-",1)
        bg=bg_cycle[i%5]
        best=bests[i] if i<len(bests) else ""
        tbl_row(pal_tbl,[num,name,desc,best],bg=bg,bold_first=False)
    doc.add_paragraph()

    h1(doc,"3.  Final Image Counts Summary")
    total_ch  = sum(s["channels"]    for s in scenes)
    total_co  = sum(s["composites"]  for s in scenes)
    total_40  = sum(s["comp40_pngs"] for s in scenes)
    total_tif = sum(s["geotiffs"]    for s in scenes)
    total_png = total_ch+total_co+total_40

    ct=doc.add_table(rows=0,cols=2)
    ct.style="Table Grid"
    hdr_row(ct,["Item","Count"])
    rows_ct=[
        ("Scenes processed","22 unique scenes (1 duplicate skipped)"),
        ("Date range","2026-06-22 to 2026-07-06"),
        ("Channel PNGs (12 × 22)",str(total_ch)),
        ("Standard composite PNGs (9 × 22)",str(total_co)),
        ("40-color composite PNGs (9×40 × 22)",str(total_40)),
        ("Total PNG images",str(total_png)),
        ("GeoTIFF files (21 × 22)",str(total_tif)),
        ("Grand total image files",str(total_png+total_tif)),
        ("Per scene: channel PNGs","12"),
        ("Per scene: standard composite PNGs","9"),
        ("Per scene: 40-color composite PNGs","9 types × 40 colors = 360"),
        ("Per scene: total PNGs","381"),
        ("Per scene: GeoTIFFs","21"),
        ("Per scene: total files","402"),
    ]
    for l,r in rows_ct:
        tbl_row(ct,[l,r],bg="EBF3FB")

    path=os.path.join(BASE,"updated_workflow.docx")
    doc.save(path)
    print(f"Saved: updated_workflow.docx")
    return path


# ── MAIN ─────────────────────────────────────────────────────────────────────
def main():
    print("Collecting statistics from output_v2/ ...")
    scenes = collect_stats()

    total_png = sum(s["total_pngs"] for s in scenes)
    total_all = sum(s["total_files"] for s in scenes)
    print(f"  {len(scenes)} scenes | {total_png} PNGs | {total_all} total files")

    print("\nGenerating per_file_summary.docx ...")
    make_per_file_summary(scenes)

    print("Generating updated_workflow.docx ...")
    make_updated_workflow(scenes)

    print("\nAll done.")
    print(f"  per_file_summary.docx  -> {BASE}")
    print(f"  updated_workflow.docx  -> {BASE}")


if __name__=="__main__":
    main()
