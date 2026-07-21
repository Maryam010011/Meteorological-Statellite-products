"""
generate_workflow.py
====================
Generates a comprehensive workflow document explaining the entire
MSG SEVIRI IODC processing pipeline from raw data to final outputs.
Saved as: workflow_document.docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

# ---------------------------------------------------------------------------
# HELPERS
# ---------------------------------------------------------------------------

def shade_cell(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def set_margins(doc, top=2.5, bottom=2.5, left=3, right=2.5):
    for s in doc.sections:
        s.top_margin = Cm(top); s.bottom_margin = Cm(bottom)
        s.left_margin = Cm(left); s.right_margin = Cm(right)

def heading(doc, text, level=1, color=(0,56,107)):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def para(doc, text, size=10.5, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor(*color)
    return p

def bullet(doc, text, size=10):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def numbered(doc, text, size=10):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text); r.font.size = Pt(size)
    return p

def code_block(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Courier New"; r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
    p.paragraph_format.left_indent = Cm(1)
    return p

def divider(doc):
    p = doc.add_paragraph()
    p.add_run("─" * 80).font.color.rgb = RGBColor(180,180,180)
    return p

def two_col_table(doc, rows_data, headers=None, hdr_color="003869"):
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    if headers:
        hr = tbl.add_row().cells
        hr[0].text = headers[0]; hr[1].text = headers[1]
        for c in hr:
            shade_cell(c, hdr_color)
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.bold = True
                    rr.font.color.rgb = RGBColor(255,255,255)
    for left, right in rows_data:
        row = tbl.add_row().cells
        row[0].text = left; row[1].text = right
        shade_cell(row[0], "EBF3FB")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold = True
    return tbl

# ---------------------------------------------------------------------------
# MAIN DOCUMENT BUILD
# ---------------------------------------------------------------------------

def build():
    doc = Document()
    set_margins(doc)

    # ========================================================================
    # COVER PAGE
    # ========================================================================
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("MSG SEVIRI IODC Satellite Data Processing")
    r.bold = True; r.font.size = Pt(22)
    r.font.color.rgb = RGBColor(0, 56, 107)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sub_p.add_run("Complete End-to-End Workflow Documentation")
    r2.bold = True; r2.font.size = Pt(16)
    r2.font.color.rgb = RGBColor(0, 112, 192)

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta_p.add_run(
        "NASTAP Assignment 4  |  Meteosat-9 IODC  |  Sub-satellite: 45.5°E\n"
        "Instrument: SEVIRI Level 1.5  |  23 .nat files  |  22 unique scenes\n"
        "Processing Date: July 2026"
    )
    r3.font.size = Pt(11); r3.italic = True
    r3.font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()
    doc.add_paragraph()

    abstract_tbl = doc.add_table(rows=1, cols=1)
    abstract_tbl.style = "Table Grid"
    ac = abstract_tbl.rows[0].cells[0]
    shade_cell(ac, "EBF3FB")
    ac.paragraphs[0].add_run("DOCUMENT ABSTRACT").bold = True
    ac.paragraphs[0].runs[0].font.size = Pt(10)
    ab = ac.add_paragraph(
        "This document provides a complete workflow description of the MSG SEVIRI IODC "
        "satellite data processing pipeline developed for NASTAP Assignment 4. It covers "
        "every stage from raw Level 1.5 native (.nat) file ingestion through environment "
        "setup, data inspection, image generation, quality correction, and final report "
        "production. The workflow processed 22 unique Meteosat-9 IODC scenes acquired "
        "between 22 June 2026 and 6 July 2026, generating 462 georeferenced PNG images "
        "and 462 GeoTIFF files, plus 21 interpretation reports and 21 user guides."
    )
    for rr in ab.runs: rr.font.size = Pt(10)

    doc.add_page_break()

    # ========================================================================
    # TABLE OF CONTENTS (manual)
    # ========================================================================
    heading(doc, "Table of Contents", 1)
    toc_items = [
        ("1", "Project Overview and Objectives"),
        ("2", "Data Description — MSG SEVIRI IODC"),
        ("3", "Software Environment Setup"),
        ("4", "Step 1 — Data Inspection (inspect_msg.py)"),
        ("5", "Step 2 — Visualization Pipeline Design"),
        ("6", "Step 3 — Geometry and Projection Corrections"),
        ("7", "Step 4 — Color and Enhancement Corrections"),
        ("8", "Step 5 — Image Rendering and Metadata Labeling"),
        ("9", "Step 6 — Batch Processing All 22 Scenes"),
        ("10", "Step 7 — Composite Generation (Separate Pass)"),
        ("11", "Step 8 — Report and User Guide Generation"),
        ("12", "Output Structure and File Inventory"),
        ("13", "Technical Challenges and Solutions"),
        ("14", "Key Scripts Reference"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.add_run("{}. {}".format(num, title)).font.size = Pt(10)
        p.paragraph_format.left_indent = Cm(0.5)

    doc.add_page_break()
    return doc

def section_1(doc):
    heading(doc, "1.  Project Overview and Objectives", 1)
    para(doc,
        "This project was carried out as part of NASTAP Assignment 4. The objective was to "
        "take raw EUMETSAT Meteosat Second Generation (MSG) SEVIRI Level 1.5 native satellite "
        "data files and produce a complete set of scientifically calibrated, visually correct, "
        "and professionally labeled meteorological image products — matching the quality and "
        "appearance of EUMETSAT's own operational viewer (view.eumetsat.int).")

    doc.add_paragraph()
    heading(doc, "Objectives", 2, (0,112,192))
    for obj in [
        "Inspect and validate all 23 .nat files — detect duplicates, read metadata.",
        "Generate 12 individual SEVIRI channel images per scene (VIS, NIR, IR, WV, HRV).",
        "Generate 9 standard RGB composite products per scene using Satpy's built-in "
        "EUMETSAT-calibrated enhancement recipes.",
        "Apply correct geostationary-to-PlateCarree reprojection with proper pillow-square "
        "disk shape, white background, yellow coastlines.",
        "Label every image with product name, satellite, sub-satellite longitude, and "
        "sensing UTC time parsed from the .nat file header.",
        "Export all products as PNG (visualization) and GeoTIFF (analysis-ready georeferenced).",
        "Write per-scene manifests (JSON + CSV) and a global summary.",
        "Produce 21 interpretation reports and 21 two-page user guides in .docx format.",
    ]:
        bullet(doc, obj)

    doc.add_paragraph()
    heading(doc, "Final Counts", 2, (0,112,192))
    two_col_table(doc, [
        ("Raw .nat files provided",        "23 files (258.6 MB total)"),
        ("Unique scenes (after dedup)",    "22 scenes"),
        ("Duplicate detected and skipped", "msg15 (15).nat = msg15 (14).nat (MD5 match)"),
        ("Date range",                     "2026-06-22 00:00 UTC  to  2026-07-06 03:45 UTC"),
        ("Channel PNGs generated",         "264  (12 channels × 22 scenes)"),
        ("Composite PNGs generated",       "198  (9 composites × 22 scenes)"),
        ("Total PNGs",                     "462"),
        ("Total GeoTIFFs",                 "462"),
        ("Interpretation Reports",         "21 .docx files"),
        ("User Guides",                    "21 .docx files"),
    ], headers=["Item", "Value"])
    doc.add_paragraph()

def section_2(doc):
    heading(doc, "2.  Data Description — MSG SEVIRI IODC", 1)
    para(doc,
        "The Meteosat Second Generation (MSG) satellite series is operated by EUMETSAT. "
        "Meteosat-9 is positioned at 45.5°E in geostationary orbit as the Indian Ocean "
        "Data Coverage (IODC) service satellite. It carries the Spinning Enhanced Visible "
        "and Infrared Imager (SEVIRI) instrument.")

    doc.add_paragraph()
    heading(doc, "SEVIRI Instrument Characteristics", 2, (0,112,192))
    two_col_table(doc, [
        ("Instrument",          "SEVIRI (Spinning Enhanced Visible and InfraRed Imager)"),
        ("Satellite",           "Meteosat-9"),
        ("Orbital position",    "45.5°E (IODC — Indian Ocean Data Coverage)"),
        ("Scan repeat cycle",   "15 minutes (full disk)"),
        ("Spatial resolution",  "3 km at nadir for 11 channels; 1 km for HRV"),
        ("Spectral channels",   "12 channels: VIS006, VIS008, IR016, IR039, WV062, WV073, "
                                "IR087, IR097, IR108, IR120, IR134, HRV"),
        ("Data format",         "EUMETSAT Level 1.5 Native (.nat) — compressed binary"),
        ("Geographic coverage", "Full Earth disk: ~80°S–80°N, ~35°W–165°E (from 45.5°E)"),
    ], headers=["Parameter", "Value"])

    doc.add_paragraph()
    heading(doc, "Data Files Provided", 2, (0,112,192))
    para(doc,
        "23 .nat files were provided in the data/ folder with filenames msg15.nat, "
        "msg15 (1).nat through msg15 (22).nat. These filenames do NOT follow EUMETSAT's "
        "standard naming convention (which encodes the sensing time), requiring a "
        "workaround to enable Satpy's reader.")

    para(doc,
        "Sensing times span four calendar dates across 22 unique scenes:",
        bold=True)
    for date_info in [
        "2026-06-22: 14 scenes (00:00 to 07:30 UTC) — night through morning transition",
        "2026-06-23: 1 scene (01:00 UTC)",
        "2026-06-29: 3 scenes (00:15 to 01:00 UTC)",
        "2026-07-06: 4 scenes (00:45 to 03:45 UTC)",
    ]:
        bullet(doc, date_info)
    doc.add_paragraph()

def section_3(doc):
    heading(doc, "3.  Software Environment Setup", 1)
    para(doc,
        "The entire processing pipeline runs on Python 3.12 under Windows. "
        "All packages were installed into the system Python environment at "
        "C:\\Users\\User\\AppData\\Local\\Programs\\Python\\Python312\\. "
        "Note: the system PATH was shadowed by an MSYS2 Python 3.14 installation, "
        "requiring all script invocations to use the full interpreter path.")

    doc.add_paragraph()
    heading(doc, "Installed Package Versions", 2, (0,112,192))
    two_col_table(doc, [
        ("Python",        "3.12"),
        ("satpy",         "0.60.0  — core satellite data reading and compositing"),
        ("pyresample",    "1.35.0  — geographic resampling between projections"),
        ("numpy",         "2.5.1   — array operations"),
        ("xarray",        "2026.7.0 — labelled array framework used by Satpy"),
        ("dask",          "2026.7.1 — lazy parallel computation for large arrays"),
        ("matplotlib",    "3.11.0  — image rendering and colorbar generation"),
        ("cartopy",       "0.25.0  — geographic map projections and coastline overlays"),
        ("rasterio",      "1.4.3   — GeoTIFF export (installed during workflow)"),
        ("scipy",         "1.18.0  — array zoom for composite resampling"),
        ("shapely",       "2.1.2   — geometric operations (Cartopy dependency)"),
        ("pyproj",        "3.7.2   — coordinate reference system conversions"),
        ("python-docx",   "1.2.0   — report and guide generation"),
        ("pillow",        "12.3.0  — image I/O support"),
        ("trollimage",    "1.28.0  — Satpy image enhancement"),
        ("pyspectral",    "0.14.3  — spectral response functions for composites"),
        ("pyorbital",     "1.12.1  — orbital geometry calculations"),
    ], headers=["Package", "Version and Role"])

    doc.add_paragraph()
    heading(doc, "Key Installation Command", 2, (0,112,192))
    code_block(doc, 'pip install rasterio==1.4.3  # only missing package, installed during Step 1')
    doc.add_paragraph()


def section_4(doc):
    heading(doc, "4.  Step 1 — Data Inspection (inspect_msg.py)", 1)
    para(doc,
        "Before any processing, all 23 .nat files were inspected to validate "
        "their contents, detect duplicates, and extract metadata. This was "
        "performed by inspect_msg.py.")

    doc.add_paragraph()
    heading(doc, "What inspect_msg.py Does", 2, (0,112,192))
    for item in [
        "Pass 1 — MD5 checksums: Computes MD5 hash of every .nat file to detect "
        "byte-identical duplicates without loading the data.",
        "Duplicate detection: Compares all MD5 hashes. Found: msg15 (15).nat is an "
        "exact duplicate of msg15 (14).nat.",
        "Pass 2 — Satpy metadata: Creates a temporary hard-link with a standard "
        "EUMETSAT filename pattern so Satpy's seviri_l1b_native reader can open the file. "
        "Reads platform name, sensing start/end time, sub-satellite longitude, and "
        "available channel list.",
        "Output: Prints a summary table with filename, size, MD5, platform, sensing times, "
        "and status (OK / DUPLICATE / ERROR) for each file.",
        "SatDump check: Confirms SatDump is NOT installed — Satpy is the sole processing tool.",
    ]:
        numbered(doc, item)

    doc.add_paragraph()
    heading(doc, "The Filename Workaround", 2, (0,112,192))
    para(doc,
        "EUMETSAT .nat files are named with a specific pattern encoding the sensing time "
        "and satellite identifier, e.g.: MSG3-SEVI-MSG15-0100-NA-20260622000000.000000000Z-....nat. "
        "The provided files are named msg15 (1).nat etc., which Satpy's reader rejects. "
        "The solution: create a temporary hard-link with a valid EUMETSAT-format name "
        "pointing to the original file, load the scene, then delete the hard-link. "
        "The original files are NEVER renamed or modified.")
    code_block(doc,
        "os.link('data/msg15 (1).nat',\n"
        "        'data/MSG3-SEVI-MSG15-0100-NA-20000101000000.000000000Z-TMPIDX0001.nat')\n"
        "scene = Scene(reader='seviri_l1b_native', filenames=[tmp_path])\n"
        "# ... process ...\n"
        "os.remove(tmp_path)  # always cleaned up in finally block")
    doc.add_paragraph()

def section_5(doc):
    heading(doc, "5.  Step 2 — Visualization Pipeline Design", 1)
    para(doc,
        "The visualization script visualize_msg.py (later replaced by visualize_msg_v2.py) "
        "implements a complete per-file processing pipeline. The design principles were:")

    for item in [
        "Process one .nat file at a time to manage memory (each scene ~12 MB in memory after loading).",
        "Load each channel/composite in a fresh Scene object to prevent memory buildup.",
        "Use Satpy's seviri_l1b_native reader exclusively — the confirmed working reader for these files.",
        "Generate both PNG (for visualization) and GeoTIFF (for GIS analysis) for every product.",
        "Write per-scene manifest.json and manifest.csv recording every product's status and file path.",
        "Skip duplicate files by MD5 hash — never process the same data twice.",
        "Log all activity with timestamps to batch_log.txt for monitoring.",
    ]:
        bullet(doc, item)

    doc.add_paragraph()
    heading(doc, "Processing Flow Per Scene", 2, (0,112,192))
    for step in [
        "Create hard-link with EUMETSAT-format filename.",
        "Open Scene, load reference channel (IR_108) to read metadata (platform, sensing times).",
        "Create output directory: output_v2/<YYYYMMDDTHHmmZ>/channels/, composites/, geotiff/.",
        "For each of 12 channels: load → resample to target area → extract numpy array → "
        "render PNG with colorbar + metadata → save GeoTIFF.",
        "For each of 9 composites: load on native scene → force Dask compute → "
        "scipy zoom to 2048×2048 → render RGB PNG → save GeoTIFF.",
        "Write manifest.json (full metadata) and manifest.csv (tabular summary).",
        "Delete temporary hard-link.",
    ]:
        numbered(doc, step)
    doc.add_paragraph()


def section_6(doc):
    heading(doc, "6.  Step 3 — Geometry and Projection Corrections", 1)
    para(doc,
        "The initial version of visualize_msg.py produced distorted, unrecognisable images "
        "('squashed pillow' with warped coastlines). The root cause and fix are documented here.")

    doc.add_paragraph()
    heading(doc, "Root Cause: Wrong Projection and Extent", 2, (0,112,192))
    para(doc, "Three compounding bugs were identified:", bold=True)
    two_col_table(doc, [
        ("Bug 1 — Wrong area definition",
         "Used +proj=eqc (PlateCarree) as the resampling target but with metre-based extent "
         "(-8.9M to +8.9M) that did not match any correct geographic coverage."),
        ("Bug 2 — Wrong Cartopy extent",
         "Passed the metre-based extent to ax.imshow() with transform=ccrs.PlateCarree(), "
         "but PlateCarree expects degree extents. Cartopy placed the image at ±8.9 million "
         "degrees — completely outside the map."),
        ("Bug 3 — Forced square grid",
         "Used 2048×2048 on a domain that is not square in degrees, distorting continent shapes."),
    ], headers=["Bug", "Description"])

    doc.add_paragraph()
    heading(doc, "The Fix", 2, (0,112,192))
    para(doc,
        "The correct approach, verified empirically by testing multiple approaches:")
    for item in [
        "Resample from native geostationary grid to PlateCarree (eqc) centred on 45.5°E, "
        "extent ±8,000,000 metres, 2048×2048 pixels.",
        "Render using ccrs.PlateCarree(central_longitude=45.5) as the Cartopy CRS.",
        "Set ax.set_extent([lon_min, lon_max, -72, 72], crs=ccrs.PlateCarree()) where "
        "lon_min=45.5-72=-26.5 and lon_max=45.5+72=117.5 degrees.",
        "Pass imshow extent as degrees [lon_min, lon_max, -72, 72] with "
        "transform=ccrs.PlateCarree() (lon_0=0, not 45.5).",
        "Use plt.subplots() not plt.figure()+add_subplot() — critical for correct rendering.",
        "Remove bbox_inches='tight' from savefig — this was cropping the map to just the "
        "colorbar (producing a 5KB blank 'bar' image).",
    ]:
        numbered(doc, item)

    doc.add_paragraph()
    heading(doc, "Result", 2, (0,112,192))
    para(doc,
        "The corrected geometry produces the characteristic 'pillow square' shape of a "
        "geostationary full-disk satellite image in PlateCarree projection: the Earth disk "
        "appears as a rounded square, with black corners (space), white figure background, "
        "yellow coastlines correctly aligned with continent boundaries.")
    doc.add_paragraph()

def section_7(doc):
    heading(doc, "7.  Step 4 — Color and Enhancement Corrections", 1)
    para(doc,
        "The first batch of composite images (airmass, dust, ash, convection etc.) "
        "appeared as oversaturated neon purple/green/magenta — completely wrong colors "
        "that did not match EUMETSAT's official viewer output. This was a critical "
        "defect requiring root-cause analysis.")

    doc.add_paragraph()
    heading(doc, "Root Cause: Manual Re-stretch Overriding EUMETSAT Enhancement", 2, (0,112,192))
    para(doc,
        "The original code extracted the raw composite array and applied a custom "
        "2nd–98th percentile stretch:")
    code_block(doc,
        "arr_min = np.nanpercentile(arr, 2)\n"
        "arr_max = np.nanpercentile(arr, 98)\n"
        "arr = (arr - arr_min) / (arr_max - arr_min)  # WRONG — destroys EUMETSAT colors")
    para(doc,
        "This discarded Satpy's built-in EUMETSAT-calibrated enhancement YAML entirely. "
        "Each composite has a precisely defined color recipe (specific channel differences, "
        "gamma corrections, and display ranges) stored in Satpy's enhancement YAML files "
        "(satpy/etc/enhancements/generic.yaml and msg-specific composites). "
        "The percentile stretch ignored all of this, producing random per-scene colors.")

    doc.add_paragraph()
    heading(doc, "The Fix: Use Satpy get_enhanced_image()", 2, (0,112,192))
    para(doc,
        "The correct approach uses Satpy's built-in enhancement pipeline:")
    code_block(doc,
        "from satpy.enhancements.enhancer import get_enhanced_image\n\n"
        "# Load composite on native scene\n"
        "sc = Scene(reader='seviri_l1b_native', filenames=[tmp])\n"
        "sc.load(['dust'])\n\n"
        "# Apply EUMETSAT calibrated enhancement — NO manual stretch\n"
        "ximg = get_enhanced_image(sc['dust'])\n"
        "data = ximg.data.compute().values  # force Dask compute while file open\n"
        "arr  = np.transpose(data[:3], (1,2,0)).astype(np.float32)\n"
        "arr  = np.clip(arr, 0, 1)  # only clip, never re-stretch")
    para(doc,
        "This produces exactly the colors EUMETSAT's viewer shows — the enhancement "
        "YAML defines the scientifically correct display range and gamma for each product.")
    doc.add_paragraph()


def section_8(doc):
    heading(doc, "8.  Step 5 — Image Rendering and Metadata Labeling", 2)
    para(doc,
        "Every output PNG is rendered using matplotlib with Cartopy geographic overlays. "
        "The rendering pipeline is the same for both channels and composites.")

    doc.add_paragraph()
    heading(doc, "Rendering Pipeline", 2, (0,112,192))
    for step in [
        "Create figure: plt.subplots(figsize=(10,10), subplot_kw={'projection': "
        "ccrs.PlateCarree(central_longitude=45.5)}, facecolor='white')",
        "Set axis background to black (space outside Earth disk = black).",
        "Set geographic extent: ax.set_extent([lon_min, lon_max, -72, 72], "
        "crs=ccrs.PlateCarree())",
        "Draw satellite data: ax.imshow(arr, extent=[lon_min, lon_max, -72, 72], "
        "transform=ccrs.PlateCarree(), origin='upper', zorder=3)",
        "Add coastlines: cartopy COASTLINE feature at 50m resolution, yellow, linewidth=0.6",
        "Add country borders: cartopy BORDERS feature, yellow, linewidth=0.3",
        "Add gridlines: white semi-transparent lines at 30° intervals with lat/lon labels",
        "For channels: add horizontal colorbar with physical units (K or %)",
        "Add title with full metadata: product name | satellite | sub-sat lon | "
        "sensing start–end UTC",
        "Save: fig.savefig(path, dpi=150, facecolor='white') — NO bbox_inches='tight'",
    ]:
        numbered(doc, step)

    doc.add_paragraph()
    heading(doc, "Metadata Label Format", 2, (0,112,192))
    code_block(doc,
        "IR 10.8 um | Meteosat-9 IODC  45.5E\n"
        "2026-06-22 00:00 UTC - 00:15 UTC")
    para(doc,
        "All metadata (platform name, sensing start time, sensing end time) is parsed "
        "directly from the .nat file header via Satpy's reader — never guessed from filename.")

    doc.add_paragraph()
    heading(doc, "Output Formats", 2, (0,112,192))
    two_col_table(doc, [
        ("PNG",     "150 DPI, white background, ~700KB–1.7MB per image depending on product"),
        ("GeoTIFF", "LZW-compressed, eqc projection, float32 pixel values, "
                    "georeferenced with rasterio"),
    ], headers=["Format", "Specification"])
    doc.add_paragraph()


def section_9(doc):
    heading(doc, "9.  Step 6 — Batch Processing All 22 Scenes", 1)
    para(doc,
        "After validation on a single test scene, all 22 unique scenes were processed "
        "sequentially by batch_run.py. Sequential (not parallel) processing was chosen "
        "to manage memory and allow per-file error isolation.")

    doc.add_paragraph()
    heading(doc, "Batch Runner Design (batch_run.py)", 2, (0,112,192))
    for item in [
        "Discover all .nat files in data/ (excludes any MSG-prefixed temp files).",
        "Deduplicate by MD5 — skip msg15 (15).nat (duplicate of msg15 (14).nat).",
        "Sort 22 unique files chronologically by sensing time (reads Satpy metadata).",
        "For each file in order: call process_file() → log result → gc.collect() → next file.",
        "Write output_v2/summary.json at completion with full run statistics.",
        "All progress logged to batch_log.txt with timestamps.",
    ]:
        numbered(doc, item)

    doc.add_paragraph()
    heading(doc, "Chronological Scene Order Processed", 2, (0,112,192))
    scenes = [
        ("1", "msg15 (1).nat",  "2026-06-22 00:00 UTC"),
        ("2", "msg15 (2).nat",  "2026-06-22 00:15 UTC"),
        ("3", "msg15 (3).nat",  "2026-06-22 00:30 UTC"),
        ("4", "msg15 (4).nat",  "2026-06-22 01:00 UTC"),
        ("5", "msg15 (5).nat",  "2026-06-22 01:15 UTC"),
        ("6", "msg15 (6).nat",  "2026-06-22 01:30 UTC"),
        ("7", "msg15 (7).nat",  "2026-06-22 02:00 UTC"),
        ("8", "msg15 (8).nat",  "2026-06-22 02:30 UTC"),
        ("9", "msg15 (9).nat",  "2026-06-22 03:15 UTC"),
        ("10","msg15 (10).nat", "2026-06-22 04:00 UTC"),
        ("11","msg15 (11).nat", "2026-06-22 04:45 UTC"),
        ("12","msg15 (12).nat", "2026-06-22 05:15 UTC"),
        ("13","msg15 (13).nat", "2026-06-22 06:30 UTC"),
        ("14","msg15 (14).nat", "2026-06-22 07:30 UTC"),
        ("15","msg15.nat",      "2026-06-23 01:00 UTC"),
        ("16","msg15 (16).nat", "2026-06-29 00:15 UTC"),
        ("17","msg15 (17).nat", "2026-06-29 00:45 UTC"),
        ("18","msg15 (18).nat", "2026-06-29 01:00 UTC"),
        ("19","msg15 (19).nat", "2026-07-06 00:45 UTC"),
        ("20","msg15 (20).nat", "2026-07-06 01:45 UTC"),
        ("21","msg15 (21).nat", "2026-07-06 02:30 UTC"),
        ("22","msg15 (22).nat", "2026-07-06 03:45 UTC"),
    ]
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    for c, t in zip(tbl.rows[0].cells, ["#","Filename","Sensing Time (UTC)"]):
        c.text = t; shade_cell(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)
    for num, fname, ts in scenes:
        row = tbl.add_row().cells
        row[0].text=num; row[1].text=fname; row[2].text=ts
    doc.add_paragraph()

def section_10(doc):
    heading(doc, "10.  Step 7 — Composite Generation (Separate Pass)", 1)
    para(doc,
        "During the initial batch run, all 9 composite products failed with the error: "
        "'composite_key missing after resample (avail: [])'. "
        "This required diagnosis and a separate composites-only run.")

    doc.add_paragraph()
    heading(doc, "Root Cause: Satpy Lazy Composite Generation", 2, (0,112,192))
    para(doc,
        "When sc.load(['dust']) is called, Satpy loads the underlying source channels "
        "(IR_087, IR_108, IR_120) but the composite DataArray itself is only computed "
        "lazily — it is NOT materialised in the scene. After sc.resample(), Satpy "
        "attempts to generate the composite on the resampled grid, but "
        "available_dataset_names() returns empty because composites are not listed "
        "as 'raw datasets'.")

    doc.add_paragraph()
    heading(doc, "The Fix: Native Scene Access + Force Compute", 2, (0,112,192))
    para(doc, "The working solution:")
    code_block(doc,
        "sc = Scene(reader='seviri_l1b_native', filenames=[tmp])\n"
        "sc.load(['dust'])               # loads source channels + composite recipe\n"
        "_ = sc['dust']                  # ACCESS the composite — forces lazy eval\n"
        "ximg = get_enhanced_image(sc['dust'])\n"
        "data = ximg.data.compute().values  # FORCE Dask compute WHILE file open\n"
        "# Now safe to close scene and delete temp file\n"
        "del sc; gc.collect()\n\n"
        "# Composite is on native 3712x3712 geos grid\n"
        "# Resize to 2048x2048 using scipy.ndimage.zoom\n"
        "from scipy.ndimage import zoom\n"
        "arr = np.stack([zoom(arr[:,:,b], 2048/3712, order=1) for b in range(3)], axis=2)")

    doc.add_paragraph()
    heading(doc, "run_composites.py", 2, (0,112,192))
    para(doc,
        "A separate script run_composites.py was written to fill in missing composites "
        "without touching or deleting existing channel images. It:")
    for item in [
        "Reads the existing output_v2/summary.json to match each timestamp folder to its .nat file.",
        "Skips any composite PNG that already exists and has size > 10KB.",
        "Processes all 9 composites for each of the 22 scenes.",
        "Result: 198 composites generated, 0 errors.",
    ]:
        bullet(doc, item)
    doc.add_paragraph()


def section_11(doc):
    heading(doc, "11.  Step 8 — Report and User Guide Generation", 1)
    para(doc,
        "After all images were verified, two types of documents were generated "
        "programmatically using python-docx for all 21 meteorological products.")

    doc.add_paragraph()
    heading(doc, "Interpretation Reports (reports/*.docx)", 2, (0,112,192))
    para(doc, "Each report contains 6 sections:")
    two_col_table(doc, [
        ("Section 1", "Product Overview — description of the channel/composite"),
        ("Section 2", "Physical Basis and Calibration — what the sensor measures, units"),
        ("Section 3", "Image Interpretation Guide — color/tone key explaining what you see"),
        ("Section 4", "Sector Applications — detailed table for Agriculture, Aviation, "
                      "Natural Resource Monitoring, Natural Disaster Monitoring"),
        ("Section 5", "Known Limitations — day/night restrictions, resolution limits, etc."),
        ("Section 6", "References — EUMETSAT manuals, peer-reviewed literature"),
    ], headers=["Section", "Content"])

    doc.add_paragraph()
    heading(doc, "Two-Page User Guides (guides/*.docx)", 2, (0,112,192))
    para(doc, "Each guide is designed to fit on two pages:")
    two_col_table(doc, [
        ("Page 1 — Top",    "Header bar, product title and wavelength"),
        ("Page 1 — Middle", "Product overview box with physical basis summary"),
        ("Page 1 — Middle", "How to Read This Image table — color/tone to meaning mapping"),
        ("Page 1 — Bottom", "Sector Quick-Reference table (3 columns: Sector, Key Use, Watch For)"),
        ("Page 2 — Top",    "Detailed Applications for all 4 sectors"),
        ("Page 2 — Bottom", "Do's and Don'ts table, document footer"),
    ], headers=["Location", "Content"])

    doc.add_paragraph()
    heading(doc, "Products Covered (21 total)", 2, (0,112,192))
    products = [
        "VIS 0.6 µm (VIS006)", "VIS 0.8 µm (VIS008)", "NIR 1.6 µm (IR_016)",
        "IR 3.9 µm (IR_039)",  "WV 6.2 µm (WV_062)",  "WV 7.3 µm (WV_073)",
        "IR 8.7 µm (IR_087)",  "IR 9.7 µm (IR_097)",  "IR 10.8 µm (IR_108)",
        "IR 12.0 µm (IR_120)", "IR 13.4 µm (IR_134)", "HRV 0.7 µm",
        "Natural Color RGB",   "Airmass RGB",           "Dust RGB",
        "Ash RGB",             "Convection RGB",        "Fog RGB",
        "Night Microphysics RGB", "Day Severe Storms RGB", "IR 10.8 Colorized",
    ]
    for prod in products:
        bullet(doc, prod)
    doc.add_paragraph()

def section_12(doc):
    heading(doc, "12.  Output Structure and File Inventory", 1)
    para(doc, "The complete output directory structure is as follows:")
    code_block(doc,
        "NASTAP/Assignment 4/\n"
        "├── data/                          <- 23 original .nat files (read-only, never modified)\n"
        "│   ├── msg15.nat\n"
        "│   ├── msg15 (1).nat  ...  msg15 (22).nat\n"
        "│\n"
        "├── output_v2/                     <- All generated images\n"
        "│   ├── summary.json               <- Global run summary\n"
        "│   ├── 20260622T0000Z/\n"
        "│   │   ├── channels/              <- 12 channel PNGs\n"
        "│   │   │   ├── VIS006.png\n"
        "│   │   │   ├── IR_108.png  ... (12 files)\n"
        "│   │   ├── composites/            <- 9 composite PNGs\n"
        "│   │   │   ├── natural_color.png\n"
        "│   │   │   ├── dust_rgb.png  ... (9 files)\n"
        "│   │   ├── geotiff/               <- 21 GeoTIFF files\n"
        "│   │   │   ├── VIS006.tif  ...  dust_rgb.tif  ... (21 files)\n"
        "│   │   ├── manifest.json          <- Per-scene product manifest\n"
        "│   │   └── manifest.csv\n"
        "│   ├── 20260622T0015Z/  ... (21 more timestamp folders)\n"
        "│\n"
        "├── reports/                       <- 21 interpretation report .docx files\n"
        "│   ├── vis006_report.docx  ...  colorized_ir_clouds_report.docx\n"
        "│\n"
        "├── guides/                        <- 21 two-page user guide .docx files\n"
        "│   ├── vis006_guide.docx  ...  colorized_ir_clouds_guide.docx\n"
        "│\n"
        "├── inspect_msg.py                 <- Data inspection script\n"
        "├── visualize_msg_v2.py            <- Corrected visualization pipeline\n"
        "├── batch_run.py                   <- Sequential batch runner (channels)\n"
        "├── run_composites.py              <- Composites-only pass\n"
        "├── generate_reports.py            <- Report and guide generator\n"
        "├── generate_workflow.py           <- This workflow document generator\n"
        "└── workflow_document.docx         <- This document")

    doc.add_paragraph()
    heading(doc, "File Counts", 2, (0,112,192))
    two_col_table(doc, [
        ("Channel PNG files",     "264  (12 channels × 22 scenes)"),
        ("Composite PNG files",   "198  (9 composites × 22 scenes)"),
        ("Total PNG files",       "462"),
        ("GeoTIFF files",         "462  (one per PNG, georeferenced)"),
        ("Per-scene manifests",   "44   (22 × manifest.json + manifest.csv)"),
        ("Global summary",        "1    (output_v2/summary.json)"),
        ("Interpretation reports","21   (.docx in reports/)"),
        ("User guides",           "21   (.docx in guides/)"),
        ("Python scripts",        "6    (.py files)"),
        ("Log files",             "2    (batch_log.txt, composites_log.txt)"),
    ], headers=["Item","Count"])
    doc.add_paragraph()


def section_13(doc):
    heading(doc, "13.  Technical Challenges and Solutions", 1)

    challenges = [
        (
            "Challenge 1: Non-standard filenames",
            "The .nat files are named msg15 (1).nat etc., which Satpy rejects. "
            "EUMETSAT's reader requires a specific filename pattern encoding satellite ID "
            "and sensing time.",
            "Create a temporary hard-link with a valid EUMETSAT-format name. "
            "Process using the link. Always delete in a finally block. "
            "Original files are never renamed or modified."
        ),
        (
            "Challenge 2: Duplicate file detection",
            "msg15 (15).nat is a byte-identical copy of msg15 (14).nat. "
            "Processing it would produce a duplicate output folder.",
            "MD5 checksum comparison at startup. Any file whose MD5 matches an "
            "already-seen file is logged and skipped."
        ),
        (
            "Challenge 3: Wrong image geometry (bar/blank output)",
            "Passing metre-based extents to imshow() with PlateCarree CRS produced "
            "blank 5KB images instead of globe images. Also, bbox_inches='tight' "
            "cropped to just the colorbar.",
            "Use degree-based extents matching the eqc resampling area. "
            "Remove bbox_inches='tight'. Use plt.subplots() not plt.figure()+add_subplot()."
        ),
        (
            "Challenge 4: Composite colors wrong (neon/oversaturated)",
            "Manual percentile stretch discarded EUMETSAT's calibrated enhancement "
            "recipes, producing random per-scene colors.",
            "Use Satpy's get_enhanced_image() which applies the built-in YAML "
            "enhancement. Only clip to [0,1] — never re-stretch."
        ),
        (
            "Challenge 5: Composites missing after resample",
            "sc.resample() returned empty available_dataset_names() for all composites "
            "because composites are generated lazily and don't appear as raw datasets.",
            "Access sc[comp_key] on the native scene (not resampled). "
            "Force Dask compute with ximg.data.compute() while the temp file is still open. "
            "Then use scipy.ndimage.zoom to resize the numpy array to target resolution."
        ),
        (
            "Challenge 6: Cartopy/Shapely LinearRing crash",
            "Cartopy 0.25 + Shapely 2.x raises GEOSException when gridline "
            "draw_labels=True is used in certain projections.",
            "Wrap savefig() in try/except. On GEOSException, clear gridliners "
            "from all axes and retry save without gridlines."
        ),
        (
            "Challenge 7: Python PATH conflict",
            "The system PATH resolved 'python' to MSYS2 Python 3.14 (in C:\\msys64) "
            "instead of Python 3.12 where all packages were installed.",
            "Use the full interpreter path: "
            "C:\\Users\\User\\AppData\\Local\\Programs\\Python\\Python312\\python.exe "
            "for all script invocations."
        ),
    ]

    for title, problem, solution in challenges:
        heading(doc, title, 2, (0,112,192))
        p_para = doc.add_paragraph()
        p_para.add_run("Problem: ").bold = True
        p_para.add_run(problem).font.size = Pt(10)
        s_para = doc.add_paragraph()
        s_para.add_run("Solution: ").bold = True
        s_para.add_run(solution).font.size = Pt(10)
        doc.add_paragraph()

def section_14(doc):
    heading(doc, "14.  Key Scripts Reference", 1)

    scripts = [
        ("inspect_msg.py",
         "Data inspection and validation",
         ["Calculates MD5 checksums for all .nat files",
          "Detects and reports duplicate files",
          "Opens each file via Satpy hard-link workaround",
          "Reads platform, sensing times, sub-satellite longitude",
          "Prints summary table and channel list",
          "Checks for SatDump installation"],
         "python inspect_msg.py data/"),

        ("visualize_msg_v2.py",
         "Core visualization pipeline (corrected version)",
         ["Defines PlateCarree eqc area (2048x2048 standard, 4096x4096 HRV)",
          "Implements hard-link workaround and MD5-based dedup",
          "process_file(): channels via Satpy resample + raw values + colorbar PNG",
          "Composites via native scene access + get_enhanced_image + scipy zoom",
          "Renders all images with correct pillow-square geometry",
          "Saves GeoTIFFs with rasterio",
          "Writes manifest.json + manifest.csv per scene"],
         "python visualize_msg_v2.py data/"),

        ("batch_run.py",
         "Sequential batch runner for all 22 scenes",
         ["Discovers and deduplicates .nat files",
          "Sorts scenes chronologically by sensing time",
          "Calls process_file() for each scene sequentially",
          "Logs all output to batch_log.txt",
          "Writes output_v2/summary.json on completion"],
         "python batch_run.py"),

        ("run_composites.py",
         "Composites-only pass (fills missing composites without touching channels)",
         ["Reads summary.json to match timestamps to .nat files",
          "Skips composites that already exist",
          "Uses native scene + force compute + scipy zoom approach",
          "Logs to composites_log.txt"],
         "python run_composites.py"),

        ("generate_reports.py",
         "Generates interpretation reports and user guides",
         ["Defines 21 product content blocks (description, physics, colors, applications)",
          "make_report(): 6-section .docx interpretation report per product",
          "make_guide(): 2-page .docx user guide per product",
          "Saves to reports/ and guides/ folders"],
         "python generate_reports.py"),

        ("generate_workflow.py",
         "Generates this workflow documentation document",
         ["Builds structured multi-section .docx using python-docx",
          "Documents every step, design decision, challenge and solution",
          "Includes output structure, file inventory, script reference"],
         "python generate_workflow.py"),
    ]

    for script_name, purpose, features, cmd in scripts:
        heading(doc, script_name, 2, (0,112,192))
        p_p = doc.add_paragraph()
        p_p.add_run("Purpose: ").bold = True
        p_p.add_run(purpose)
        p_p.runs[-1].font.size = Pt(10)
        para(doc, "Key features:", bold=True, size=10)
        for f in features:
            bullet(doc, f)
        para(doc, "Usage:", bold=True, size=10)
        code_block(doc, cmd)
        doc.add_paragraph()


def main():
    print("Building workflow document...")
    doc = build()

    section_1(doc)
    doc.add_page_break()
    section_2(doc)
    doc.add_page_break()
    section_3(doc)
    doc.add_page_break()
    section_4(doc)
    doc.add_page_break()
    section_5(doc)
    doc.add_page_break()
    section_6(doc)
    doc.add_page_break()
    section_7(doc)
    doc.add_page_break()
    section_8(doc)
    doc.add_page_break()
    section_9(doc)
    doc.add_page_break()
    section_10(doc)
    doc.add_page_break()
    section_11(doc)
    doc.add_page_break()
    section_12(doc)
    doc.add_page_break()
    section_13(doc)
    doc.add_page_break()
    section_14(doc)

    # Final page — summary
    doc.add_page_break()
    heading(doc, "Document Summary", 1)
    para(doc,
        "This document has described the complete end-to-end workflow for processing "
        "MSG SEVIRI IODC Level 1.5 native satellite data files into a full set of "
        "calibrated, georeferenced, and professionally labeled meteorological image products.",
        size=11)
    doc.add_paragraph()
    two_col_table(doc, [
        ("Raw data input",            "23 .nat files, Meteosat-9 IODC 45.5°E"),
        ("Unique scenes processed",   "22 (1 duplicate skipped)"),
        ("Date range",                "2026-06-22 to 2026-07-06"),
        ("Total PNG images",          "462 (264 channels + 198 composites)"),
        ("Total GeoTIFFs",            "462"),
        ("Interpretation reports",    "21 .docx files"),
        ("User guides",               "21 .docx files"),
        ("Processing environment",    "Python 3.12, Satpy 0.60.0, Cartopy 0.25.0"),
        ("Total processing time",     "~4 hours (channels ~2.5h + composites ~0.5h)"),
        ("External data used",        "None — all images derived exclusively from provided .nat files"),
    ], headers=["Item","Value"])

    out_path = os.path.join(BASE, "workflow_document.docx")
    doc.save(out_path)
    print("Saved: {}".format(out_path))


if __name__ == "__main__":
    main()
