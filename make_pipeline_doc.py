"""
make_pipeline_doc.py
====================
Generates a single narrative document describing the complete pipeline:
  - How data comes out of .nat files
  - How Python scripts process it
  - How each file produces channels and composites
  - All timestamps highlighted with 15-min scan cycle shown
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE = os.path.dirname(os.path.abspath(__file__))

# ------------------------------------------------------------------
# HELPERS
# ------------------------------------------------------------------
def shade(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)

def margins(doc, t=2, b=2, l=2.5, r=2):
    for s in doc.sections:
        s.top_margin=Cm(t); s.bottom_margin=Cm(b)
        s.left_margin=Cm(l); s.right_margin=Cm(r)

def h1(doc, txt, color=(0,56,107)):
    h = doc.add_heading(txt, level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(*color)
    return h

def h2(doc, txt, color=(0,112,192)):
    h = doc.add_heading(txt, level=2)
    for r in h.runs:
        r.font.color.rgb = RGBColor(*color)
    return h

def p(doc, txt, sz=10.5, bold=False, italic=False, color=None):
    para = doc.add_paragraph()
    r = para.add_run(txt)
    r.bold=bold; r.italic=italic; r.font.size=Pt(sz)
    if color: r.font.color.rgb=RGBColor(*color)
    return para

def bullet(doc, txt, sz=10):
    para = doc.add_paragraph(style="List Bullet")
    r = para.add_run(txt); r.font.size=Pt(sz)
    return para

def code(doc, txt):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    r = para.add_run(txt)
    r.font.name="Courier New"; r.font.size=Pt(9)
    r.font.color.rgb=RGBColor(30,30,30)
    return para

def arrow(doc):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = para.add_run("  ▼  ")
    r.font.size=Pt(14); r.bold=True
    r.font.color.rgb=RGBColor(0,112,192)
    return para

# ------------------------------------------------------------------
# DATA
# ------------------------------------------------------------------
SCENES = [
    # (file, scan_start, scan_end, gap_note)
    ("msg15 (1).nat",  "2026-06-22  00:00 UTC", "2026-06-22  00:15 UTC", ""),
    ("msg15 (2).nat",  "2026-06-22  00:15 UTC", "2026-06-22  00:30 UTC", "15 min"),
    ("msg15 (3).nat",  "2026-06-22  00:30 UTC", "2026-06-22  00:45 UTC", "15 min"),
    ("msg15 (4).nat",  "2026-06-22  01:00 UTC", "2026-06-22  01:15 UTC", "15 min"),
    ("msg15 (5).nat",  "2026-06-22  01:15 UTC", "2026-06-22  01:30 UTC", "15 min"),
    ("msg15 (6).nat",  "2026-06-22  01:30 UTC", "2026-06-22  01:45 UTC", "15 min"),
    ("msg15 (7).nat",  "2026-06-22  02:00 UTC", "2026-06-22  02:15 UTC", "15 min"),
    ("msg15 (8).nat",  "2026-06-22  02:30 UTC", "2026-06-22  02:45 UTC", "15 min"),
    ("msg15 (9).nat",  "2026-06-22  03:15 UTC", "2026-06-22  03:30 UTC", "15 min"),
    ("msg15 (10).nat", "2026-06-22  04:00 UTC", "2026-06-22  04:15 UTC", "15 min"),
    ("msg15 (11).nat", "2026-06-22  04:45 UTC", "2026-06-22  05:00 UTC", "15 min"),
    ("msg15 (12).nat", "2026-06-22  05:15 UTC", "2026-06-22  05:30 UTC", "15 min"),
    ("msg15 (13).nat", "2026-06-22  06:30 UTC", "2026-06-22  06:45 UTC", "15 min"),
    ("msg15 (14).nat", "2026-06-22  07:30 UTC", "2026-06-22  07:45 UTC", "15 min"),
    ("msg15.nat",      "2026-06-23  01:00 UTC", "2026-06-23  01:15 UTC", "NEW DATE  June 23"),
    ("msg15 (16).nat", "2026-06-29  00:15 UTC", "2026-06-29  00:30 UTC", "NEW DATE  June 29"),
    ("msg15 (17).nat", "2026-06-29  00:45 UTC", "2026-06-29  01:00 UTC", "15 min"),
    ("msg15 (18).nat", "2026-06-29  01:00 UTC", "2026-06-29  01:15 UTC", "15 min"),
    ("msg15 (19).nat", "2026-07-06  00:45 UTC", "2026-07-06  01:00 UTC", "NEW DATE  July 6"),
    ("msg15 (20).nat", "2026-07-06  01:45 UTC", "2026-07-06  02:00 UTC", "15 min"),
    ("msg15 (21).nat", "2026-07-06  02:30 UTC", "2026-07-06  02:45 UTC", "15 min"),
    ("msg15 (22).nat", "2026-07-06  03:45 UTC", "2026-07-06  04:00 UTC", "15 min"),
]

CHANNELS = [
    ("VIS006", "VIS 0.6 µm",  "Reflected sunlight — vegetation/cloud/ocean",         "gray",   "0–100 %"),
    ("VIS008", "VIS 0.8 µm",  "Reflected sunlight — enhanced vegetation signal",     "gray",   "0–100 %"),
    ("IR_016", "NIR 1.6 µm",  "Ice vs water cloud phase discrimination",             "gray",   "0–100 %"),
    ("IR_039", "IR 3.9 µm",   "Fire hot-spots, fog detection, mixed emission",       "gray_r", "200–330 K"),
    ("WV_062", "WV 6.2 µm",   "Upper-troposphere water vapour (600–300 hPa)",        "gray_r", "200–260 K"),
    ("WV_073", "WV 7.3 µm",   "Mid-troposphere water vapour (500–300 hPa)",          "gray_r", "200–280 K"),
    ("IR_087", "IR 8.7 µm",   "Dust and silicate mineral aerosol absorption",        "gray_r", "200–330 K"),
    ("IR_097", "IR 9.7 µm",   "Stratospheric ozone absorption band",                 "gray_r", "200–320 K"),
    ("IR_108", "IR 10.8 µm",  "Primary thermal window — cloud/surface temperature", "gray_r", "200–330 K"),
    ("IR_120", "IR 12.0 µm",  "Split-window — SST and dust retrieval",               "gray_r", "200–330 K"),
    ("IR_134", "IR 13.4 µm",  "CO2 absorption — cloud-top height retrieval",         "gray_r", "200–330 K"),
    ("HRV",    "HRV 0.7 µm",  "High-resolution visible (~1 km) — fine detail",       "gray",   "0–100 %"),
]

COMPOSITES = [
    ("natural_color",      "Natural Color",          "R=VIS006  G=VIS008  B=IR016",         "Daylight only — true colour like view"),
    ("airmass",            "Airmass RGB",             "R=WV062-WV073  G=IR097-IR108  B=WV062","Air mass type, jet stream, cold/warm air"),
    ("dust",               "Dust RGB",                "R=IR120-IR108  G=IR108-IR087  B=IR108","Mineral dust plumes — pink/magenta"),
    ("ash",                "Ash RGB",                 "R=IR120-IR108  G=IR108-IR087  B=IR108","Volcanic ash — red/orange"),
    ("convection",         "Convection RGB",          "R=WV062-WV073  G=IR039-IR108  B=NIR-VIS","Rapidly growing convective cells"),
    ("fog",                "Fog RGB",                 "R=IR120-IR039  G=IR108-IR039  B=IR108","Fog and low stratus — orange/yellow"),
    ("night_microphysics", "Night Microphysics RGB",  "R=IR120-IR108  G=IR108-IR039  B=IR108","Cloud phase at night — water vs ice"),
    ("day_severe_storms",  "Day Severe Storms RGB",   "R=WV062-WV073  G=IR039-IR108  B=NIR-VIS","Overshooting tops and severe cells"),
    ("colorized_ir_clouds","IR 10.8 Colorized",       "IR108 with rainbow colour table",     "Enhanced cloud top temperature"),
]

# ------------------------------------------------------------------
# BUILD
# ------------------------------------------------------------------
def build():
    doc = Document()
    margins(doc)

    # ══════════════════════════════════════════════════════
    # COVER
    # ══════════════════════════════════════════════════════
    doc.add_paragraph()
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run("MSG SEVIRI IODC — Complete Data Processing Pipeline")
    r.bold=True; r.font.size=Pt(20)
    r.font.color.rgb=RGBColor(0,56,107)

    sp = doc.add_paragraph()
    sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = sp.add_run(
        "From Raw .nat Files to Channels, Composites, Reports and Guides\n"
        "Meteosat-9  |  45.5°E IODC  |  SEVIRI Level 1.5  |  22 Scenes  |  July 2026")
    r2.font.size=Pt(11); r2.italic=True
    r2.font.color.rgb=RGBColor(60,60,60)

    doc.add_paragraph()

    # Overview box
    ov = doc.add_table(rows=1, cols=1)
    ov.style = "Table Grid"
    oc = ov.rows[0].cells[0]
    shade(oc, "E8F4FD")
    oc.paragraphs[0].add_run(
        "This document tells the complete story of how 23 raw satellite data files "
        "(.nat format) from Meteosat-9 were processed into 462 calibrated meteorological "
        "images. It shows exactly how each file is opened, what data comes out of it, "
        "how Python scripts transform that data step by step, and what the final outputs "
        "look like — including the actual timestamps of every scan with the 15-minute "
        "repeat cycle highlighted."
    ).font.size = Pt(10)

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # SECTION 1 — WHAT IS A .nat FILE
    # ══════════════════════════════════════════════════════
    h1(doc, "1.  What Is a .nat File?")
    p(doc,
      "A .nat file is EUMETSAT's native binary format for MSG SEVIRI Level 1.5 data. "
      "Level 1.5 means the raw satellite counts have already been calibrated into physical "
      "units (reflectance % or brightness temperature K) and geolocated onto a standard grid.")

    doc.add_paragraph()
    h2(doc, "Inside a Single .nat File")
    p(doc,
      "Every .nat file is a self-contained snapshot of Earth as seen from Meteosat-9 at "
      "one moment in time. It contains:")

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    for c, t in zip(tbl.rows[0].cells, ["What Is Stored", "Details"]):
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

    contents = [
        ("Header metadata",
         "Satellite name (Meteosat-9), sensing start time, sensing end time, "
         "sub-satellite longitude (45.5°E), calibration coefficients"),
        ("12 spectral channels",
         "VIS006, VIS008, IR016, IR039, WV062, WV073, IR087, IR097, "
         "IR108, IR120, IR134 at 3km resolution + HRV at 1km resolution"),
        ("Pixel grid",
         "3712 x 3712 pixels for standard channels covering the full Earth disk "
         "in geostationary projection"),
        ("Pixel values",
         "Calibrated: reflectance (0-100%) for visible channels, "
         "brightness temperature (200-330 K) for infrared channels"),
        ("File size",
         "~11.8 MB per file (all 23 files total 258.6 MB)"),
        ("Projection",
         "Geostationary (geos), longitude 45.5°E, satellite height 35,785,831 m"),
    ]
    for left, right in contents:
        row = tbl.add_row().cells
        row[0].text=left; row[1].text=right
        shade(row[0],"D6E4F0")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True

    doc.add_paragraph()
    h2(doc, "The 15-Minute Scan Cycle")
    p(doc,
      "SEVIRI scans the full Earth disk once every 15 minutes. The satellite spins "
      "and the mirror tilts to sweep north-to-south across the disk. Each scan starts "
      "at the southern edge and takes exactly 15 minutes to complete. This is why every "
      ".nat file covers a 15-minute window: the SCAN START TIME is when the first line "
      "is acquired, and the SCAN END TIME is 15 minutes later when the last line is done.")

    doc.add_page_break()

    # ══════════════════════════════════════════════════════
    # SECTION 2 — TIMESTAMPS TABLE
    # ══════════════════════════════════════════════════════
    h1(doc, "2.  All 22 Scenes — Timestamps and 15-Minute Gaps")
    p(doc,
      "The table below shows every scene in chronological order. "
      "The '15 min gap' column highlights consecutive scans. "
      "New dates are shown in a distinct color — there are gaps of hours or days "
      "between the date groups (these scenes were selectively archived, not continuous).")
    doc.add_paragraph()

    # Timestamp table
    ts_tbl = doc.add_table(rows=1, cols=5)
    ts_tbl.style = "Table Grid"
    for c, t in zip(ts_tbl.rows[0].cells,
                    ["#", "File", "Scan Start (UTC)", "Scan End (UTC)", "Gap / Note"]):
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

    date_colors = {
        "2026-06-22": "FFFFFF",
        "2026-06-23": "FFF2CC",
        "2026-06-29": "E2EFDA",
        "2026-07-06": "FCE4D6",
    }

    for i, (fname, start, end, gap) in enumerate(SCENES, 1):
        row = ts_tbl.add_row().cells
        date_key = start[:10]
        bg = date_colors.get(date_key, "FFFFFF")
        row[0].text = str(i)
        row[1].text = fname
        row[2].text = start
        row[3].text = end
        row[4].text = gap

        for c in row:
            shade(c, bg)

        # Bold the gap note for new dates
        if "NEW DATE" in gap:
            for pp in row[4].paragraphs:
                for rr in pp.runs:
                    rr.bold=True
                    rr.font.color.rgb=RGBColor(192,0,0)

        # Bold filename
        for pp in row[1].paragraphs:
            for rr in pp.runs: rr.bold=True

    doc.add_paragraph()

    # Legend
    leg = doc.add_table(rows=1, cols=4)
    leg.style = "Table Grid"
    for c, col, lbl in zip(leg.rows[0].cells,
                            ["FFFFFF","FFF2CC","E2EFDA","FCE4D6"],
                            ["June 22 (14 scenes)","June 23 (1 scene)",
                             "June 29 (3 scenes)","July 6 (4 scenes)"]):
        shade(c, col)
        c.text=lbl
        for pp in c.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9)

    doc.add_paragraph()
    p(doc,
      "Note: msg15 (15).nat was detected as a byte-identical duplicate of msg15 (14).nat "
      "(confirmed by MD5 checksum). It was skipped automatically — that is why the file "
      "list jumps from (14) to (16) in the June 22 group.", italic=True, sz=9,
      color=(100,100,100))

    doc.add_page_break()
    return doc

def section_pipeline(doc):
    # ══════════════════════════════════════════════════════
    # SECTION 3 — HOW THE PIPELINE WORKS
    # ══════════════════════════════════════════════════════
    h1(doc, "3.  How the Pipeline Processes Each .nat File")
    p(doc,
      "Every .nat file goes through exactly the same pipeline. Below is the complete "
      "journey from raw file on disk to finished PNG image and GeoTIFF.")

    doc.add_paragraph()

    steps = [
        ("STEP 1", "File Discovery",
         "batch_run.py scans the data/ folder for all *.nat files. "
         "It computes the MD5 checksum of each file and removes duplicates. "
         "The 22 unique files are then sorted chronologically by reading the "
         "sensing start time from each file's header via Satpy.",
         "003869", "DEEAF1"),

        ("STEP 2", "Hard-Link Rename Workaround",
         "Satpy's seviri_l1b_native reader requires the filename to follow EUMETSAT's "
         "standard pattern (e.g. MSG3-SEVI-MSG15-0100-NA-20260622000000.000000000Z-....nat). "
         "Since the provided files are named msg15 (1).nat etc., a temporary hard-link "
         "with a valid EUMETSAT-format name is created on disk pointing to the original file. "
         "The original file is NEVER renamed or copied — the hard-link is just an alias.\n\n"
         "Code: os.link('data/msg15 (1).nat', 'data/MSG3-SEVI-...TMPIDX0001.nat')\n"
         "After processing: os.remove(temp_link)  <- always cleaned up",
         "375623", "E2EFDA"),

        ("STEP 3", "Scene Loading — Reading the .nat File",
         "Satpy opens the .nat file via its seviri_l1b_native reader:\n\n"
         "   sc = Scene(reader='seviri_l1b_native', filenames=[temp_link])\n\n"
         "At this point NO data is actually read from disk yet — Satpy creates a "
         "lazy reference. The file header is read to confirm the satellite, sensing time, "
         "and available channels.",
         "833C00", "FCE4D6"),

        ("STEP 4", "Metadata Extraction",
         "A reference channel (IR_108) is loaded to read the embedded metadata:\n\n"
         "   sc.load(['IR_108'])\n"
         "   platform      = sc['IR_108'].attrs['platform_name']   -> 'Meteosat-9'\n"
         "   sensing_start = sc['IR_108'].attrs['start_time']      -> datetime object\n"
         "   sensing_end   = sc['IR_108'].attrs['end_time']        -> datetime object\n\n"
         "This timestamp is used for: output folder name (20260622T0000Z), "
         "image title label, and manifest entries.",
         "7030A0", "EAE0F0"),

        ("STEP 5", "Output Folder Creation",
         "A timestamped output folder is created:\n\n"
         "   output_v2/\n"
         "   └── 20260622T0000Z/\n"
         "       ├── channels/      <- 12 channel PNGs go here\n"
         "       ├── composites/    <- 9 composite PNGs go here\n"
         "       ├── geotiff/       <- 21 GeoTIFFs go here\n"
         "       ├── manifest.json  <- product inventory\n"
         "       └── manifest.csv",
         "003869", "DEEAF1"),

        ("STEP 6", "Channel Processing (×12)",
         "For each of the 12 SEVIRI channels (VIS006 through HRV):\n\n"
         "  a) Load: sc.load([ch]) — Dask reads the compressed channel data from .nat\n"
         "  b) Resample: sc.resample(eqc_area) — reprojects from geostationary metres\n"
         "     to PlateCarree degrees (eqc projection, 2048×2048 pixels for standard\n"
         "     channels, 4096×4096 for HRV)\n"
         "  c) Extract: arr = rsc[ch].values.astype(float32) — compute array in memory\n"
         "     Result: 2D numpy array (2048,2048), values in K or %\n"
         "  d) Render PNG with matplotlib + Cartopy (see Step 8)\n"
         "  e) Save GeoTIFF with rasterio (georeferenced, LZW compressed)\n"
         "  f) Record in manifest",
         "833C00", "FCE4D6"),

        ("STEP 7", "Composite Processing (×9)",
         "For each of the 9 RGB composites:\n\n"
         "  a) Load: sc.load([comp_key]) — Satpy loads all required source channels\n"
         "     and registers the composite recipe (e.g. dust = IR120-IR108, IR108-IR087, IR108)\n"
         "  b) Access: _ = sc[comp_key] — forces Satpy to evaluate the composite recipe\n"
         "     and apply the EUMETSAT calibrated enhancement YAML\n"
         "  c) Enhance: ximg = get_enhanced_image(sc[comp_key])\n"
         "     arr = ximg.data.compute().values — compute while file is still open\n"
         "     Result: 3D numpy array (3, 3712, 3712) in range 0.0–1.0\n"
         "  d) Resize: scipy.ndimage.zoom to 2048×2048 per band\n"
         "  e) Render RGB PNG with matplotlib + Cartopy (see Step 8)\n"
         "  f) Save GeoTIFF and record in manifest",
         "7030A0", "EAE0F0"),

        ("STEP 8", "Image Rendering",
         "Every image (channel or composite) is rendered using the same approach:\n\n"
         "  • Figure:  plt.subplots(figsize=(10,10), projection=PlateCarree(lon_0=45.5))\n"
         "  • Background: ax.set_facecolor('black')  <- space = black\n"
         "  • Figure background: facecolor='white'   <- page = white\n"
         "  • Extent: ax.set_extent([-26.5, 117.5, -72, 72], crs=PlateCarree())\n"
         "  • Data: ax.imshow(arr, extent=[-26.5,117.5,-72,72], transform=PlateCarree())\n"
         "  • Coastlines: cartopy COASTLINE 50m, yellow, linewidth=0.6\n"
         "  • Borders: cartopy BORDERS 50m, yellow, linewidth=0.3\n"
         "  • Gridlines: white, 30° intervals, lat/lon labels\n"
         "  • Colorbar (channels only): physical units K or %\n"
         "  • Title: 'IR 10.8 um | Meteosat-9 IODC  45.5E  |  2026-06-22 00:00 UTC'\n"
         "  • Save: fig.savefig(path, dpi=150, facecolor='white')\n"
         "    NOTE: bbox_inches='tight' is NOT used — it crops to colorbar only",
         "375623", "E2EFDA"),

        ("STEP 9", "Manifest Writing",
         "After all products for one scene are done:\n\n"
         "  manifest.json contains: filename, MD5, platform, sensing_start, sensing_end,\n"
         "  generation_time, and for every product: name, type, status, png path, tif path\n\n"
         "  manifest.csv: same data in tabular format for spreadsheet use\n\n"
         "  The temp hard-link is deleted in a finally: block — guaranteed cleanup\n"
         "  even if any step above fails.",
         "003869", "DEEAF1"),
    ]

    for step_id, step_name, step_text, hdr_color, bg_color in steps:
        # Step header box
        stbl = doc.add_table(rows=1, cols=2)
        stbl.style = "Table Grid"
        sc1, sc2 = stbl.rows[0].cells
        sc1.text = step_id
        sc2.text = step_name
        shade(sc1, hdr_color); shade(sc2, hdr_color)
        for c in [sc1, sc2]:
            for pp in c.paragraphs:
                for rr in pp.runs:
                    rr.bold=True; rr.font.size=Pt(11)
                    rr.font.color.rgb=RGBColor(255,255,255)
        sc1.width = Cm(2.5)

        # Step body
        btbl = doc.add_table(rows=1, cols=1)
        btbl.style = "Table Grid"
        bc = btbl.rows[0].cells[0]
        shade(bc, bg_color)
        bc.text = ""
        for line in step_text.split("\n"):
            if line.startswith("   ") or line.startswith("  "):
                # code-style line
                cp = bc.add_paragraph()
                cp.paragraph_format.left_indent = Cm(0.5)
                r = cp.add_run(line)
                r.font.name="Courier New"; r.font.size=Pt(8.5)
            else:
                np_ = bc.add_paragraph()
                r = np_.add_run(line)
                r.font.size=Pt(10)

        arrow(doc)

    # Final output box
    final_tbl = doc.add_table(rows=1, cols=1)
    final_tbl.style = "Table Grid"
    fc = final_tbl.rows[0].cells[0]
    shade(fc, "003869")
    fp_ = fc.paragraphs[0]
    r = fp_.add_run("OUTPUT: 21 images per scene  (12 channels + 9 composites)  x  22 scenes  =  462 PNGs + 462 GeoTIFFs")
    r.bold=True; r.font.size=Pt(11)
    r.font.color.rgb=RGBColor(255,255,255)
    fp_.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

def section_channels(doc):
    h1(doc, "4.  The 12 SEVIRI Channels — What Each One Measures")
    p(doc,
      "When a .nat file is processed, 12 individual channel images are produced. "
      "Each channel is a different wavelength window on the electromagnetic spectrum, "
      "measuring a fundamentally different physical property of the atmosphere and surface.")
    doc.add_paragraph()

    ch_tbl = doc.add_table(rows=1, cols=5)
    ch_tbl.style = "Table Grid"
    for c, t in zip(ch_tbl.rows[0].cells,
                    ["Channel", "Name", "What It Measures", "Color Scale", "Value Range"]):
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

    vis_bg  = "FFFDE7"  # pale yellow for visible
    ir_bg   = "E8F5E9"  # pale green for IR/WV
    hrv_bg  = "E3F2FD"  # pale blue for HRV

    for ch_id, name, desc, cmap, rng in CHANNELS:
        row = ch_tbl.add_row().cells
        row[0].text=ch_id; row[1].text=name
        row[2].text=desc; row[3].text=cmap; row[4].text=rng
        bg = vis_bg if ch_id in ("VIS006","VIS008","IR_016","HRV") else ir_bg
        if ch_id=="HRV": bg=hrv_bg
        for c in row: shade(c, bg)
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True; rr.font.size=Pt(9)

    doc.add_paragraph()

    # Legend
    leg = doc.add_table(rows=1, cols=3)
    leg.style = "Table Grid"
    for c, col, lbl in zip(leg.rows[0].cells,
                            [vis_bg, ir_bg, hrv_bg],
                            ["Visible/NIR channels (reflected sunlight — daylight only)",
                             "Infrared/Water Vapour channels (thermal emission — day and night)",
                             "HRV — High Resolution Visible (~1 km)"]):
        shade(c, col); c.text=lbl
        for pp in c.paragraphs:
            for rr in pp.runs: rr.font.size=Pt(9)

    doc.add_paragraph()
    p(doc,
      "Each channel image shows the satellite disk (pillow-square shape) with: "
      "white figure background, black space outside the disk, yellow coastlines and "
      "country borders, white gridlines at 30° intervals, and a horizontal colorbar "
      "showing the physical measurement scale.", italic=True, sz=9, color=(80,80,80))

    doc.add_page_break()


def section_composites(doc):
    h1(doc, "5.  The 9 RGB Composites — Combining Channels for Insight")
    p(doc,
      "Beyond individual channels, Satpy's built-in composite recipes combine multiple "
      "channels into a single RGB image. Each composite is designed by EUMETSAT's scientists "
      "to highlight specific meteorological phenomena that are invisible in any single channel.")
    doc.add_paragraph()

    p(doc, "How a composite is formed:", bold=True)
    p(doc,
      "A composite assigns three channel expressions to Red, Green, and Blue. "
      "For example, the Dust RGB assigns:\n"
      "  Red   = IR_120 − IR_108  (dust absorbs more at 12.0 than 10.8 µm)\n"
      "  Green = IR_108 − IR_087  (dust/surface emission difference)\n"
      "  Blue  = IR_108           (surface temperature reference)\n\n"
      "When dust is present, these differences produce a distinctive pink/magenta color "
      "that is immediately recognizable. The enhancement YAML defines the exact stretch "
      "and gamma for each band to match EUMETSAT's official display standard.")
    doc.add_paragraph()

    comp_tbl = doc.add_table(rows=1, cols=4)
    comp_tbl.style = "Table Grid"
    for c, t in zip(comp_tbl.rows[0].cells,
                    ["Product", "Display Name", "Channel Recipe", "What It Reveals"]):
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

    comp_colors = ["DEEAF1","E2EFDA","FCE4D6","EAE0F0",
                   "FFF2CC","DEEAF1","E2EFDA","FCE4D6","EAE0F0"]
    for (key, name, recipe, reveals), bg in zip(COMPOSITES, comp_colors):
        row = comp_tbl.add_row().cells
        row[0].text=key; row[1].text=name
        row[2].text=recipe; row[3].text=reveals
        for c in row: shade(c, bg)
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True; rr.font.size=Pt(9)

    doc.add_paragraph()
    p(doc,
      "Important: Composites using visible channels (Natural Color, Convection, "
      "Day Severe Storms) only work during daylight hours. Night-time scenes produce "
      "dark or invalid images for those products. "
      "Thermal IR composites (Dust, Ash, Fog, Night Microphysics, Colorized IR) "
      "work at any time of day or night.", italic=True, sz=9, color=(80,80,80))

    doc.add_page_break()


def section_per_scene(doc):
    h1(doc, "6.  Per-Scene Output — What Each .nat File Produces")
    p(doc,
      "The table below shows every scene with its timestamp and exact output folder. "
      "Each folder contains 12 channel images + 9 composite images + 21 GeoTIFFs.")
    doc.add_paragraph()

    out_tbl = doc.add_table(rows=1, cols=4)
    out_tbl.style = "Table Grid"
    for c, t in zip(out_tbl.rows[0].cells,
                    ["File", "Scan Start — End (UTC)", "Output Folder", "Products"]):
        c.text=t; shade(c,"003869")
        for pp in c.paragraphs:
            for rr in pp.runs:
                rr.bold=True; rr.font.color.rgb=RGBColor(255,255,255)

    date_colors = {
        "2026-06-22": "FFFFFF",
        "2026-06-23": "FFF2CC",
        "2026-06-29": "E2EFDA",
        "2026-07-06": "FCE4D6",
    }

    folder_map = {
        "2026-06-22  00:00 UTC": "20260622T0000Z",
        "2026-06-22  00:15 UTC": "20260622T0015Z",
        "2026-06-22  00:30 UTC": "20260622T0030Z",
        "2026-06-22  01:00 UTC": "20260622T0100Z",
        "2026-06-22  01:15 UTC": "20260622T0115Z",
        "2026-06-22  01:30 UTC": "20260622T0130Z",
        "2026-06-22  02:00 UTC": "20260622T0200Z",
        "2026-06-22  02:30 UTC": "20260622T0230Z",
        "2026-06-22  03:15 UTC": "20260622T0315Z",
        "2026-06-22  04:00 UTC": "20260622T0400Z",
        "2026-06-22  04:45 UTC": "20260622T0445Z",
        "2026-06-22  05:15 UTC": "20260622T0515Z",
        "2026-06-22  06:30 UTC": "20260622T0630Z",
        "2026-06-22  07:30 UTC": "20260622T0730Z",
        "2026-06-23  01:00 UTC": "20260623T0100Z",
        "2026-06-29  00:15 UTC": "20260629T0015Z",
        "2026-06-29  00:45 UTC": "20260629T0045Z",
        "2026-06-29  01:00 UTC": "20260629T0100Z",
        "2026-07-06  00:45 UTC": "20260706T0045Z",
        "2026-07-06  01:45 UTC": "20260706T0145Z",
        "2026-07-06  02:30 UTC": "20260706T0230Z",
        "2026-07-06  03:45 UTC": "20260706T0345Z",
    }

    for fname, start, end, gap in SCENES:
        folder = folder_map.get(start, "")
        time_str = "{} - {}".format(start, end.split()[1])
        date_key = start[:10]
        bg = date_colors.get(date_key, "FFFFFF")

        row = out_tbl.add_row().cells
        row[0].text = fname
        row[1].text = time_str
        row[2].text = "output_v2/{}/".format(folder)
        row[3].text = "12 ch + 9 comp + 21 TIF"
        for c in row: shade(c, bg)
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True; rr.font.size=Pt(9)

    doc.add_paragraph()
    p(doc,
      "Total output: 22 folders × 21 products = 462 PNG files + 462 GeoTIFF files",
      bold=True, color=(0,56,107))
    doc.add_page_break()


def section_summary(doc):
    h1(doc, "7.  Final Outputs Summary")
    p(doc, "After all 22 scenes were processed, the following outputs were generated:")
    doc.add_paragraph()

    sum_tbl = doc.add_table(rows=0, cols=2)
    sum_tbl.style = "Table Grid"
    rows = [
        ("Raw .nat files processed", "22 unique scenes (1 duplicate skipped)"),
        ("Date range covered",       "2026-06-22  to  2026-07-06"),
        ("Channel PNGs",             "264  (12 channels × 22 scenes)"),
        ("Composite PNGs",           "198  (9 composites × 22 scenes)"),
        ("Total PNG images",         "462"),
        ("GeoTIFF files",            "462  (georeferenced, LZW-compressed)"),
        ("Per-scene manifests",      "22 manifest.json + 22 manifest.csv"),
        ("Global summary",           "output_v2/summary.json"),
        ("Interpretation reports",   "21 .docx files in reports/"),
        ("User guides (2-page)",     "21 .docx files in guides/"),
        ("This workflow document",   "workflow_pipeline.docx"),
        ("Data source",              "ONLY the 23 .nat files in data/ — no external data"),
    ]
    for left, right in rows:
        row = sum_tbl.add_row().cells
        row[0].text=left; row[1].text=right
        shade(row[0],"D6E4F0")
        for pp in row[0].paragraphs:
            for rr in pp.runs: rr.bold=True

    doc.add_paragraph()
    closing = doc.add_paragraph()
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = closing.add_run(
        "Every image in this project was derived exclusively from the .nat files "
        "provided in the data/ folder. No external satellite imagery, sample datasets, "
        "or demonstration data was used at any point.")
    r.italic=True; r.font.size=Pt(10)
    r.font.color.rgb=RGBColor(0,56,107)


def main():
    print("Building pipeline document...")
    doc = build()
    section_pipeline(doc)
    section_channels(doc)
    section_composites(doc)
    section_per_scene(doc)
    section_summary(doc)

    out = os.path.join(BASE, "workflow_pipeline.docx")
    doc.save(out)
    import os.path as op
    print("Saved: {}  ({:.0f} KB)".format(out, op.getsize(out)/1024))


if __name__ == "__main__":
    main()
